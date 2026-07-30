from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "workspace-ui-specification.md"
DOCX_PATH = BASE_DIR / "workspace-ui-specification.docx"
DIAGRAM_DIR = BASE_DIR / "diagrams" / "ui"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
HEADER_FILL = "E8EEF5"
BORDER = "A7B3C2"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill="#F2F4F7", outline="#A7B3C2"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    fnt = font(17)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 24)
    total_h = len(lines) * 21
    y = y1 + max(8, ((y2 - y1 - total_h) // 2))
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill="#0B2545", font=fnt)
        y += 21


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#41546A"):
    sx, sy = start
    ex, ey = end
    draw.line([start, end], fill=color, width=3)
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 12 * direction, ey - 7), (ex - 12 * direction, ey + 7)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 7, ey - 12 * direction), (ex + 7, ey - 12 * direction)]
    draw.polygon(pts, fill=color)


def parse_mermaid_node(token: str) -> tuple[str, str]:
    token = token.strip()
    match = re.match(r"^([A-Za-z][A-Za-z0-9_]*)(?:\[(?:\"([^\"]+)\"|'([^']+)'|([^\]]+))\]|\{(?:\"([^\"]+)\"|'([^']+)'|([^}]+))\})?$", token)
    if not match:
        return token, token
    node_id = match.group(1)
    label = next((g for g in match.groups()[1:] if g), node_id)
    return node_id, label


def render_flowchart(lines: list[str], title: str, name: str) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    nodes: dict[str, str] = {}
    edges: list[tuple[str, str, str]] = []

    for raw in lines[1:]:
        line = raw.strip()
        if not line or "-->" not in line:
            continue
        left, right = line.split("-->", 1)
        edge_label = ""
        right = right.strip()
        if right.startswith("|"):
            parts = right.split("|", 2)
            if len(parts) == 3:
                edge_label = parts[1].strip()
                right = parts[2].strip()
        src_id, src_label = parse_mermaid_node(left)
        dst_id, dst_label = parse_mermaid_node(right)
        nodes.setdefault(src_id, src_label)
        nodes.setdefault(dst_id, dst_label)
        edges.append((src_id, dst_id, edge_label))

    if not nodes:
        path = DIAGRAM_DIR / f"{name}.png"
        img = Image.new("RGB", (1200, 260), "white")
        draw = ImageDraw.Draw(img)
        draw.text((40, 40), title, fill="#1F4D78", font=font(28, bold=True))
        draw.text((40, 110), "\n".join(lines), fill="#0B2545", font=font(16))
        img.save(path)
        return path

    inbound = {node: 0 for node in nodes}
    outgoing: dict[str, list[str]] = {node: [] for node in nodes}
    for src, dst, _ in edges:
        inbound[dst] = inbound.get(dst, 0) + 1
        outgoing.setdefault(src, []).append(dst)

    sources = [node for node, count in inbound.items() if count == 0] or [next(iter(nodes))]
    levels = {node: 0 for node in sources}
    queue = list(sources)
    while queue:
        src = queue.pop(0)
        for dst in outgoing.get(src, []):
            next_level = levels[src] + 1
            if dst not in levels or next_level > levels[dst]:
                levels[dst] = next_level
                queue.append(dst)
    for node in nodes:
        levels.setdefault(node, 0)

    by_level: dict[int, list[str]] = {}
    for node, level in levels.items():
        by_level.setdefault(level, []).append(node)

    max_level = max(by_level)
    max_rows = max(len(items) for items in by_level.values())
    vertical_layout = max_level >= 6
    width = 1200 if vertical_layout else 1400
    height = max(520, 130 + (max_level + 1) * 105) if vertical_layout else max(430, 130 + max_rows * 108)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 28), title, fill="#1F4D78", font=font(28, bold=True))

    positions: dict[str, tuple[int, int, int, int]] = {}
    box_h = 72
    if vertical_layout:
        for level, items in by_level.items():
            items = sorted(items)
            count = len(items)
            gap = 18
            box_w = int(min(250, max(145, (width - 120 - (count - 1) * gap) / max(1, count))))
            total_w = count * box_w + (count - 1) * gap
            start_x = (width - total_w) // 2
            y = 95 + level * 105
            for row, node in enumerate(items):
                x = start_x + row * (box_w + gap)
                positions[node] = (x, y, x + box_w, y + box_h)
    else:
        box_w = 200
        x_gap = (width - 120 - box_w) / max(1, max_level)
        for level, items in by_level.items():
            items = sorted(items)
            col_x = 60 + int(level * x_gap)
            total_h = len(items) * box_h + (len(items) - 1) * 34
            start_y = 105 + max(0, (height - 135 - total_h) // 2)
            for row, node in enumerate(items):
                y = start_y + row * (box_h + 34)
                positions[node] = (col_x, y, col_x + box_w, y + box_h)

    for src, dst, label in edges:
        if src not in positions or dst not in positions:
            continue
        sx1, sy1, sx2, sy2 = positions[src]
        dx1, dy1, dx2, dy2 = positions[dst]
        start = (sx2, sy1 + (sy2 - sy1) // 2)
        end = (dx1, dy1 + (dy2 - dy1) // 2)
        if dx1 <= sx1:
            start = (sx1 + (sx2 - sx1) // 2, sy2)
            end = (dx1 + (dx2 - dx1) // 2, dy1)
        arrow(draw, start, end)
        if label:
            mx = (start[0] + end[0]) // 2
            my = (start[1] + end[1]) // 2 - 14
            draw.rounded_rectangle((mx - 42, my - 12, mx + 42, my + 12), radius=8, fill="white", outline="#D9E2EC")
            bbox = draw.textbbox((0, 0), label, font=font(12))
            draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 8), label, fill="#41546A", font=font(12))

    fills = ["#DDEBF7", "#E8EEF5", "#F2F4F7", "#FFF7E6", "#E7F4E4"]
    for node, xy in positions.items():
        draw_box(draw, xy, nodes[node], fill=fills[levels[node] % len(fills)])

    path = DIAGRAM_DIR / f"{name}.png"
    img.save(path)
    return path


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margin: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag in ("top", "bottom", "start", "end"):
        elem = mar.find(qn(f"w:{tag}"))
        if elem is None:
            elem = OxmlElement(f"w:{tag}")
            mar.append(elem)
        elem.set(qn("w:w"), str(80 if tag in ("top", "bottom") else margin))
        elem.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    set_table_borders(table)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "WarpTalk - Workspace UI SRS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(95, 95, 95)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - Workspace UI Software Requirement Specification"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(95, 95, 95)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Workspace UI Software Requirement Specification")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("WarpTalk Web Module - SRS for Workspace Screens and RBAC Flows")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(BLUE)


def add_markdown_table(doc: Document, rows: list[str]):
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) < 2:
        return
    headers = parsed[0]
    body = parsed[2:] if len(parsed) > 1 and all(set(c.replace(":", "").replace("-", "").strip()) == set() for c in parsed[1]) else parsed[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [max(900, 9360 // len(headers)) for _ in headers]
    diff = 9360 - sum(widths)
    widths[-1] += diff
    set_table_geometry(table, widths)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(INK)

    for row in body:
        cells = table.add_row().cells
        values = row[: len(headers)] + [""] * max(0, len(headers) - len(row))
        for i, value in enumerate(values[: len(headers)]):
            cells[i].text = value.replace("`", "")
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_diagram(doc: Document, title: str, image_path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run(title)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(95, 95, 95)


def add_paragraph_with_bold_label(doc: Document, line: str):
    if line.startswith("**") and "**:" in line:
        label, value = line.split("**:", 1)
        label = label.strip("*")
        p = doc.add_paragraph()
        r = p.add_run(label + ": ")
        r.bold = True
        p.add_run(value.strip())
        return True
    return False


def build_docx():
    source = MD_PATH.read_text(encoding="utf-8").splitlines()
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    for old_diagram in DIAGRAM_DIR.glob("*.png"):
        old_diagram.unlink()
    doc = Document()
    setup_styles(doc)
    add_title(doc)

    i = 0
    diagram_counter = 0
    current_section = "Workspace UI"
    while i < len(source):
        line = source[i].rstrip()
        if not line:
            i += 1
            continue
        if line == "---":
            i += 1
            continue

        if line.startswith("```"):
            fence_type = line.strip().strip("`").strip()
            fence_lines: list[str] = []
            i += 1
            while i < len(source) and not source[i].startswith("```"):
                fence_lines.append(source[i])
                i += 1
            if fence_type == "mermaid" and fence_lines and fence_lines[0].strip().startswith("flowchart"):
                title = f"{current_section} - Screen flow"
                diagram_counter += 1
                safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", title.lower()).strip("_")[:70] or "flowchart"
                safe_name = f"{diagram_counter:02d}_{safe_name}"
                image_path = render_flowchart(fence_lines, title, safe_name)
                add_diagram(doc, title, image_path)
            else:
                add_code_block(doc, fence_lines)
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(source) and source[i].startswith("|"):
                table_lines.append(source[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if line.startswith("# "):
            # Title already rendered as a cover title.
            i += 1
            continue
        if line.startswith("## "):
            current_section = line[3:].strip()
            doc.add_heading(current_section, level=1)
            i += 1
            continue
        if line.startswith("### "):
            heading = line[4:].strip()
            if heading.lower() != "screen flow":
                current_section = heading
            doc.add_heading(heading, level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip().replace("`", ""))
            i += 1
            continue
        if len(line) > 3 and line[0].isdigit() and ". " in line[:5]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(line.split(". ", 1)[1].strip().replace("`", ""))
            i += 1
            continue

        if not add_paragraph_with_bold_label(doc, line):
            doc.add_paragraph(line.replace("`", ""))
        i += 1

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
