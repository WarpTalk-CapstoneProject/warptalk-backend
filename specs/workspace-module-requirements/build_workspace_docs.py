from __future__ import annotations

import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
DIAGRAM_DIR = BASE_DIR / "diagrams"
MD_PATH = BASE_DIR / "workspace-module-overview.md"
DOCX_PATH = BASE_DIR / "workspace-software-requirement-specification.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "A7B3C2"


SOURCE_SPECS = [
    ("WT-139", "../139-workspace-creation-selection/spec.md", "Workspace creation, listing, selection, active context"),
    ("WT-139-AC", "../139-workspace-creation-selection/workspace-types-and-role-permissions-acceptance-criteria.md", "Enterprise workspace role, membership type and permission matrix"),
    ("WT-140", "../140-workspace-invitations/spec.md", "Invitation lifecycle and email/token rules"),
    ("WT-141", "../141-workspace-members/spec.md", "Member listing, role changes, soft-delete, ownership protection"),
    ("WT-157", "../157-workspace-enterprise-external-collaboration/spec.md", "Enterprise verified domains and external collaborator isolation"),
    ("WT-158", "../158-workspace-document-access-policy/spec.md", "Document metadata, ACL precedence, external access boundary"),
    ("WT-158-Approval", "../158-workspace-document-access-policy/spec-approval-workflow.md", "Document ingestion approval workflow"),
    ("WT-158-AI", "../158-workspace-document-access-policy/spec-ai-guardrails.md", "AI guardrails and policy inheritance"),
    ("WT-158-Logic", "../158-workspace-document-access-policy/handled-document-logic.md", "Document handling and access decision logic"),
    ("WT-159", "../159-workspace-govern-internal-meetings-artifacts/spec.md", "Meeting governance, artifact retention, gRPC boundary"),
    ("DB/AI Guardrails", "../feat-2026-06-03-update-db-for-build-library-and-ai-guardrails.md", "Document library database and AI context boundary"),
    ("Local Encryption", "../feat-2026-06-07-local-document-encryption-aes256.md", "Local storage encryption using workspace-derived keys"),
    ("Identity Enrichment", "../refactor-2026-06-04-workspace-identity-enrichment-approach-1.md", "Auth identity snapshots and service boundary"),
    ("PII Presidio TD", "../techdebt-2026-06-07-pii-presidio-api-transition.md", "Future PII scanner transition to Presidio"),
    ("System Spec Reference", "reference-google-doc.txt", "Document control, technology stack, limitations, QC checklist, DoD pattern"),
    ("UI Mainflow Source", "https://docs.google.com/document/d/1xObm3bnGcMPOx71I2u-XC4VdNG886pvJlyj7TFshLAQ/edit?tab=t.0", "Workspace governance screen behavior, UI rules, loading/empty/error/success states and cross-page rules"),
    ("Workspace UI Spec", "workspace-ui-specification.md", "Separated UI specification for Workspace screens based on UI Mainflow Source, not warptalk-web implementation"),
    ("RabbitMQ Official", "https://www.rabbitmq.com/docs", "Messaging workflow, exchanges, queues, consumers, acknowledgements, publisher confirms and dead-letter exchanges"),
    ("Infrastructure DB", "../../../warptalk-infrastructure/scripts/init-db.sql", "Physical PostgreSQL schema and workspace foreign keys"),
    ("ERD Guidelines - Lucidchart", "https://www.lucidchart.com/pages/er-diagrams", "ERD concepts: entities, attributes, keys, relationships, cardinality, physical model guidance"),
    ("ERD Syntax - Mermaid", "https://mermaid.js.org/syntax/entityRelationshipDiagram.html", "Crow's-foot ERD notation used for Markdown diagram source"),
    ("Workspace Unit/Integration Tests", "../../workspace/tests/WarpTalk.WorkspaceService.Tests", "Implemented xUnit/NSubstitute/Testcontainers coverage for Workspace service, controllers, ACL, ingestion and middleware."),
    ("Backend Postman Collections", "../../test/postman", "Backend-level manual/E2E API collections and environment used for API smoke/regression verification."),
]


REQUIREMENTS = [
    ("FR-WS-001", "Workspace", "Người dùng đã xác thực có thể tạo Enterprise Workspace và được gán Owner trong cùng transaction.", "WT-139"),
    ("FR-WS-002", "Workspace", "Hệ thống MUST NOT tự tạo workspace cá nhân mặc định hoặc phân nhánh workspace type; workspace hiện tại là Enterprise tenant boundary.", "WT-139-AC"),
    ("FR-WS-003", "Workspace", "Danh sách workspace phải có phân trang, tìm kiếm và chỉ trả về workspace mà user là thành viên active.", "WT-139"),
    ("FR-WS-004", "Workspace", "Người dùng phải chọn active workspace trước khi dùng room, transcript, billing hoặc tài nguyên workspace.", "WT-139, WT-139-AC"),
    ("FR-WS-005", "Security", "Gateway/Auth phải truyền workspace context nội bộ có chữ ký cho downstream services.", "WT-139"),
    ("FR-WS-006", "Invitation", "Owner/Admin chỉ được mời thành viên vào Enterprise Workspace theo verified-domain và external-collaboration policy.", "WT-140"),
    ("FR-WS-007", "Invitation", "Invitation có các trạng thái pending, accepted, revoked, expired, replaced; resend thay thế token cũ.", "WT-140"),
    ("FR-WS-008", "Invitation", "Accept invitation yêu cầu email đăng nhập khớp tuyệt đối với email được mời.", "WT-140"),
    ("FR-WS-009", "Member", "Member list hỗ trợ phân trang/tìm kiếm và chỉ hiện member active, với hạn chế riêng cho External Member.", "WT-141, WT-157"),
    ("FR-WS-010", "Member", "Remove member là soft-delete, lưu RemovedAt/RemovedBy và không xóa lịch sử audit/billing/meeting.", "WT-141"),
    ("FR-WS-011", "Member", "Workspace luôn phải còn ít nhất một active Owner; chặn owner cuối cùng rời, bị remove hoặc bị demote.", "WT-141"),
    ("FR-WS-012", "Enterprise", "A user can belong to many Enterprise Workspaces, but can be Internal in at most one domain-verified Enterprise Workspace; other cross-organization memberships must be External.", "WT-157"),
    ("FR-WS-013", "Enterprise", "Email ngoài verified domain chỉ được mời khi AllowExternalCollaboration=true và bị ép role External Member.", "WT-157"),
    ("FR-WS-014", "Enterprise", "External Member không được quản trị workspace và chỉ xem tài nguyên meeting mà họ là participant.", "WT-157, WT-159"),
    ("FR-WS-015", "Document", "Workspace Document Library lưu metadata, storage pointer, document type, AI flags, retention và audit data.", "WT-158, DB/AI Guardrails"),
    ("FR-WS-016", "Document", "Document ACL dùng deny-overrides; explicit deny thắng allow, sensitive document deny-by-default.", "WT-158"),
    ("FR-WS-017", "Document", "Document pending/awaiting approval chỉ Owner/Admin/Uploader/Document Owner được truy cập.", "WT-158-Approval, WT-158-AI"),
    ("FR-WS-018", "Document", "Tài liệu bị archive/delete không được dùng làm context cho AI/RAG.", "WT-158, DB/AI Guardrails"),
    ("FR-WS-019", "Document", "Upload chỉ chấp nhận PDF, DOCX, TXT và meeting artifacts theo chính sách kích thước/extension.", "WT-158"),
    ("FR-WS-020", "Security", "Local storage provider phải mã hóa AES-256-CBC và xác thực HMAC-SHA512 với key dẫn xuất theo workspace.", "Local Encryption"),
    ("FR-WS-021", "AI", "Document ingestion phát sự kiện qua Redis Stream kết hợp RabbitMQ, phân loại sensitive/AI eligibility và audit kết quả.", "WT-158-AI, PII Presidio TD, RabbitMQ Official"),
    ("FR-WS-022", "Meeting", "TranslationRoomService kiểm tra member/policy workspace qua gRPC, không cross-join DB.", "WT-159"),
    ("FR-WS-023", "Meeting", "Workspace policy kiểm soát max active rooms, allowed target languages và artifact retention.", "WT-159"),
    ("FR-WS-024", "Glossary", "Workspace-level glossary hỗ trợ business_domain để AI/translation dùng đúng ngữ cảnh phòng ban.", "DB/AI Guardrails"),
    ("FR-WS-025", "Workspace", "Cung cấp API lấy số liệu thống kê tổng hợp của Workspace (Dashboard Stats) bao gồm thành viên, tài liệu, glossary, credit và phòng họp.", "Specs + Workspace/TranslationRoom integration"),
    ("FR-WS-026", "Workspace", "Cung cấp API lấy danh sách nhật ký hoạt động (Dashboard Activities) hỗ trợ tìm kiếm, phân trang và lọc theo hành động.", "Specs + Workspace audit"),
    ("FR-WS-027", "Meeting governance", "Owner/Admin có thể bật/tắt quyền tạo meeting cho từng workspace member bằng cột `workspace_members.can_create_meetings`; TranslationRoom phải validate cờ này qua Workspace boundary trước khi tạo room.", "WT-159 design decision"),
]

DETAILED_FR_SPECS = {
    "FR-WS-001": (
        "Cho phép người dùng đã xác thực (JWT hợp lệ) gửi yêu cầu tạo một không gian làm việc doanh nghiệp mới.",
        "POST /api/v1/workspaces",
        "Xác thực thông tin đầu vào (tên workspace, domain doanh nghiệp); Tạo bản ghi Workspace mới với UUID v7 và sinh slug duy nhất (xử lý trùng lặp tự động bằng suffix); Thêm bản ghi membership liên kết user với vai trò Owner (Internal Member); Lưu tất cả thay đổi trong một transaction (Unit of Work)."
    ),
    "FR-WS-002": (
        "Hệ thống hoạt động theo mô hình cô lập dữ liệu B2B Enterprise duy nhất. Không tự động provision các loại workspace cá nhân (Personal/Default Workspace) cho tài khoản mới.",
        "Internal Provision Logic",
        "Mọi thao tác sử dụng phòng họp hay tài liệu phải chạy trên ngữ cảnh Enterprise Workspace được tạo hoặc gia nhập rõ ràng; Tránh cung cấp bất kỳ API hoặc luồng ngầm nào để provision không gian làm việc cá nhân cho tài khoản mới."
    ),
    "FR-WS-003": (
        "Trả về danh sách các Enterprise Workspace mà tài khoản hiện tại đang là thành viên hoạt động.",
        "GET /api/v1/workspaces",
        "Truy vấn danh sách các workspace của người dùng có `RemovedAt IS NULL`; Hỗ trợ lọc theo tên workspace và phân trang kết quả trả về."
    ),
    "FR-WS-004": (
        "Người dùng chọn một workspace làm ngữ cảnh hoạt động hiện tại.",
        "POST /api/v1/workspaces/{id}/select",
        "Xác thực thành viên hoạt động trong workspace; Đọc vai trò và cấu hình chi tiết của workspace; Lưu ngữ cảnh làm việc hiện tại vào cache Redis/session."
    ),
    "FR-WS-005": (
        "API Gateway hoặc Middleware xác thực tự động đính kèm header chữ ký nội bộ chứa WorkspaceId và UserId đã ký mật mã vào mọi request truyền xuống các downstream services.",
        "Internal Middleware Context Service",
        "Bảo vệ chống giả mạo ngữ cảnh (context spoofing); Downstream services giải mã và xác thực chữ ký trước khi xử lý."
    ),
    "FR-WS-006": (
        "Cho phép Owner hoặc Admin gửi lời mời tham gia workspace đến một địa chỉ email.",
        "POST /api/v1/workspaces/{workspaceId}/invitations",
        "Kiểm tra email thuộc verified domain hay email ngoài; Email ngoài domain chỉ được mời khi AllowExternalCollaboration=true và bị ép vai trò External Member; Sinh token lời mời duy nhất, mã hóa hash lưu DB; Sau khi commit invitation, hệ thống gọi Email/Notification service để gửi email trực tiếp tới receiver bằng địa chỉ email được mời, trong email chỉ chứa raw invite token/link một lần và không expose token hash."
    ),
    "FR-WS-007": (
        "Quản lý các trạng thái và vòng đời của lời mời.",
        "GET/DELETE /api/v1/workspaces/{workspaceId}/invitations",
        "Các trạng thái hợp lệ gồm Pending, Accepted, Revoked, Expired, và Replaced; Gửi lại lời mời cho cùng email sẽ tự động hủy token cũ (chuyển sang Replaced) và cấp token mới."
    ),
    "FR-WS-008": (
        "Người dùng chấp nhận lời mời bằng cách cung cấp token.",
        "POST /api/v1/workspaces/invitations/accept",
        "Xác thực token còn hiệu lực (Pending và chưa hết hạn); Đăng nhập bằng tài khoản có email trùng khớp tuyệt đối (case-insensitive) với email được mời; Tạo membership mới và chuyển trạng thái invite sang Accepted."
    ),
    "FR-WS-009": (
        "Trả về danh sách thành viên đang hoạt động trong workspace.",
        "GET /api/v1/workspaces/{workspaceId}/members",
        "Hỗ trợ phân trang và tìm kiếm theo tên/email; Chặn External Member không được phép xem danh sách thành viên nội bộ."
    ),
    "FR-WS-010": (
        "Cho phép Owner hoặc Admin loại bỏ một thành viên ra khỏi workspace.",
        "DELETE /api/v1/workspaces/{workspaceId}/members/{userId}",
        "Bản ghi membership được cập nhật removed_at và removed_by (soft delete); Không xóa vật lý dữ liệu để đảm bảo tính toàn vẹn của lịch sử cuộc họp và audit logs."
    ),
    "FR-WS-011": (
        "Hệ thống ngăn chặn việc tự rời khỏi workspace, hạ cấp vai trò, hoặc xóa thành viên của người dùng là Owner duy nhất còn hoạt động.",
        "POST /api/v1/workspaces/{workspaceId}/members/transfer-ownership",
        "Yêu cầu phải chuyển nhượng quyền sở hữu (Transfer Ownership) sang thành viên khác trước khi rời hoặc demote."
    ),
    "FR-WS-012": (
        "WarpTalk remains multi-workspace: one account can belong to many Enterprise Workspaces, but can be Internal in at most one domain-verified Enterprise Workspace.",
        "Internal Domain Enforcement Logic",
        "Backend create/accept-invite flows reject a second Internal Home Workspace; External memberships remain allowed when workspace policy permits them."
    ),
    "FR-WS-013": (
        "Kiểm soát chính sách mời thành viên bên ngoài hệ thống dựa trên cấu hình cho phép.",
        "Workspace Settings Logic",
        "Nếu AllowExternalCollaboration là false, chặn hoàn toàn việc mời email ngoài verified domain; Nếu bật, email ngoài domain chỉ được mời với vai trò External Member."
    ),
    "FR-WS-014": (
        "Cô lập và hạn chế quyền hạn của thành viên bên ngoài để bảo vệ tài nguyên doanh nghiệp.",
        "External Member Access Boundary",
        "Chặn truy cập thư mục thành viên nội bộ, cài đặt workspace, và toàn bộ tài liệu dùng chung; Chỉ cho phép truy cập tài nguyên phòng họp được mời trực tiếp trong grace period được cấu hình."
    ),
    "FR-WS-015": (
        "Quản lý thư viện tài liệu tri thức dùng làm ngữ cảnh cho AI.",
        "POST/GET /api/v1/workspaces/{workspaceId}/documents",
        "Bản ghi tài liệu lưu trữ metadata, storage key/provider, cờ is_sensitive, và trạng thái AI ingestion."
    ),
    "FR-WS-016": (
        "Áp dụng quy tắc kiểm tra quyền truy cập tài liệu tri thức.",
        "Document Access Evaluator API",
        "Explicit deny ghi đè allow (deny-overrides); Tài liệu nhạy cảm (is_sensitive=true) mặc định bị chặn truy cập đối với Member thường trừ khi là uploader hoặc quản trị viên."
    ),
    "FR-WS-017": (
        "Quy trình phê duyệt tài liệu tải lên bởi thành viên thông thường trước khi ingestion.",
        "POST /api/v1/workspaces/{workspaceId}/documents/{documentId}/approve",
        "Tài liệu do Member tải lên đi vào trạng thái AwaitingApproval; Owner/Admin duyệt để kích hoạt hoặc từ chối để xóa mềm tài liệu."
    ),
    "FR-WS-018": (
        "Khi tài liệu bị xóa mềm hoặc lưu trữ (archived), hệ thống lập tức loại bỏ các chunks/vector liên quan ra khỏi Vector DB.",
        "RabbitMQ Invalidation Event Flow",
        "Thay đổi cờ `ai_eligible = false` trong DB; Gửi sự kiện qua RabbitMQ đến AI worker để tiến hành xóa vector points nhằm tránh rò rỉ dữ liệu."
    ),
    "FR-WS-019": (
        "Chặn tải lên các file không đúng chính sách định dạng hoặc kích thước.",
        "Upload Document API Validation",
        "Chỉ chấp nhận các định dạng PDF, DOCX, TXT; Kiểm tra và chặn các file vượt quá kích thước giới hạn tối đa quy định trong settings."
    ),
    "FR-WS-020": (
        "Mã hóa dữ liệu nhạy cảm đối với storage cục bộ.",
        "Local Storage Cryptography Provider",
        "Sử dụng thuật toán AES-256-CBC để mã hóa file binary và HMAC-SHA512 để xác thực tính toàn vẹn; Khóa mã hóa được sinh động học từ workspace context."
    ),
    "FR-WS-021": (
        "Tích hợp RabbitMQ để truyền sự kiện tải lên, phê duyệt và xóa tài liệu bất đồng bộ đến worker xử lý AI.",
        "RabbitMQ Event Publisher & Consumer",
        "Sử dụng publisher confirms ở phía gửi và manual acknowledgements ở phía nhận; Triển khai hàng đợi retry (TTL) và hàng đợi lỗi Dead-Letter Queue (DLQ) để đảm bảo độ tin cậy hệ thống."
    ),
    "FR-WS-022": (
        "TranslationRoom Service gọi gRPC ngược về Workspace Service để kiểm tra tính hợp lệ của thành viên và cấu hình trước khi tạo/join phòng họp.",
        "Workspace gRPC Policy Endpoint",
        "Xác thực vai trò member, cờ active, và kiểm tra các chính sách về room của workspace."
    ),
    "FR-WS-023": (
        "Áp dụng cấu hình room limit, allowed target languages và artifact retention period.",
        "gRPC Policy Retrieval",
        "gRPC service đọc cấu hình JSONB settings của workspace và trả về cho room/artifact service thực thi."
    ),
    "FR-WS-024": (
        "Quản lý thuật ngữ glossary chuyên ngành phục vụ dịch thuật AI theo từng phòng ban.",
        "POST/GET/DELETE /api/v1/workspaces/{workspaceId}/documents/glossaries",
        "Hỗ trợ phân tách glossary theo business domain và ngôn ngữ nguồn/đích; Đảm bảo không trùng lặp thuật ngữ trong cùng domain."
    ),
    "FR-WS-025": (
        "API tổng hợp các chỉ số hoạt động của Workspace phục vụ dashboard quản trị.",
        "GET /api/v1/workspaces/{id}/dashboard/stats",
        "Truy vấn DB lấy thống kê thành viên, tài liệu và glossary; Gọi gRPC đến TranslationRoom Service để lấy số phòng họp và tổng số phút dịch thuật; Lưu kết quả vào Redis Cache tối đa 5 phút."
    ),
    "FR-WS-026": (
        "Trả về danh sách nhật ký hoạt động để phục vụ kiểm toán (audit trail).",
        "GET /api/v1/workspaces/{id}/dashboard/activities",
        "Hỗ trợ phân trang, tìm kiếm và lọc theo loại hoạt động; Chỉ hiển thị các thay đổi cấu hình, thành viên và thao tác tài liệu nhạy cảm."
    ),
    "FR-WS-027": (
        "Cho phép Owner/Admin cấp hoặc thu hồi quyền tạo meeting cho từng member mà không đổi role của member đó.",
        "PATCH /api/v1/workspaces/{workspaceId}/members/{userId}/meeting-permission",
        "Persist vào `workspace.workspace_members.can_create_meetings`; Internal Member mặc định true, External Member mặc định false; removed member luôn bị chặn dù flag còn true."
    )
}

FUNCTIONAL_REQUIREMENTS_SCOPE = (
    "Functional requirements chỉ mô tả hành vi hệ thống có thể quan sát hoặc kiểm thử được. "
    "Các requirement dưới đây được tổng hợp từ WT-139/140/141/157/158/159 và đối chiếu với code hiện tại của Workspace Service "
    "ở API controllers, Application services, Domain entities/rules và Infrastructure repositories/cache/event worker. "
    "Các thuộc tính chất lượng như bảo mật, hiệu năng, availability, accessibility được tách riêng ở phần Non-functional Requirements."
)


FUNCTIONAL_SOURCE_DETAILS = {
    "FR-WS-001": "WT-139; code: WorkspacesController.CreateWorkspace, WorkspaceService.CreateWorkspaceAsync, WorkspaceRepository, WorkspaceMemberRepository",
    "FR-WS-002": "WT-139-AC; code: không có WorkspaceType enum/column, không có personal workspace auto-provision flow",
    "FR-WS-003": "WT-139; code: WorkspacesController.GetWorkspaces, WorkspaceService.GetWorkspacesAsync, WorkspaceRepository.GetWorkspacesForUserAsync",
    "FR-WS-004": "WT-139/AC; code: WorkspacesController.SelectWorkspace, WorkspaceService.SelectWorkspaceAsync, WorkspaceCacheService Redis",
    "FR-WS-005": "WT-139; code/interface: active workspace context, Redis cache/session, downstream signed context contract",
    "FR-WS-006": "WT-140/157; code: WorkspaceInvitationsController, WorkspaceInvitationService.InviteMemberAsync",
    "FR-WS-007": "WT-140; code: WorkspaceInvitation entity/status enum, Preview/Accept/Revoke/List invitation flows",
    "FR-WS-008": "WT-140; code: WorkspaceInvitationService.AcceptInvitationAsync email/token validation",
    "FR-WS-009": "WT-141/157; code: WorkspaceMembersController.GetMembers, WorkspaceMemberService.ListMembersAsync",
    "FR-WS-010": "WT-141; code: WorkspaceMemberService.RemoveMemberAsync, WorkspaceMember.RemovedAt/RemovedBy",
    "FR-WS-011": "WT-141; code: CountActiveOwnersAsync, ChangeMemberRoleAsync, TransferOwnershipAsync",
    "FR-WS-012": "WT-157; code: WorkspaceHelper.IsUserInternalMemberOfAnyEnterpriseWorkspaceAsync, verified-domain repositories",
    "FR-WS-013": "WT-157; code: WorkspaceInvitationService external policy checks, UpdateWorkspaceSettingsAsync",
    "FR-WS-014": "WT-157/159; code: MembershipType.External checks, ListMembers external guard, DocumentAccessEvaluator",
    "FR-WS-015": "WT-158; code: WorkspaceDocumentsController.UploadDocument, WorkspaceDocumentService.UploadDocumentAsync",
    "FR-WS-016": "WT-158; code: DocumentAccessEvaluator.EvaluateAccessAsync, WorkspaceDocumentAccessPolicyRepository",
    "FR-WS-017": "WT-158-Approval/AI; code: WorkspaceDocumentService.ApproveDocumentAsync",
    "FR-WS-018": "WT-158/DB-AI; code: DeleteDocumentAsync, RedisDocumentEventPublisher delete/archive events",
    "FR-WS-019": "WT-158; code: UploadDocumentRequest, WorkspaceDocumentHelper.GenerateStorageKey, WorkspaceDocument metadata",
    "FR-WS-020": "Local Encryption spec; code-boundary: local storage provider requirement, workspace-derived crypto metadata",
    "FR-WS-021": "WT-158-AI/PII/RabbitMQ; code: IWorkspaceDocumentEventPublisher, RedisDocumentEventPublisher, DocumentAiIngestionConsumerService; target: RabbitMQ durable delivery",
    "FR-WS-022": "WT-159; code/interface: TranslationRoom client boundary, no cross-service DB join",
    "FR-WS-023": "WT-159; code/spec: WorkspaceSettingsDto/WorkspaceConfiguration, room/artifact policy contract",
    "FR-WS-024": "DB/AI Guardrails; schema: workspace_knowledge_glossaries, AI/translation prompt adapter boundary",
    "FR-WS-025": "Specs; code: WorkspaceDashboardController.GetStats, WorkspaceDashboardService, TranslationRoomGrpcClient",
    "FR-WS-026": "Specs; code: WorkspaceDashboardController.GetActivities, WorkspaceDashboardService, WorkspaceDocumentAuditRepository",
    "FR-WS-027": "WT-159 decision; target code: WorkspaceMember.CanCreateMeetings, WorkspaceMemberService toggle, TranslationRoom workspace validation boundary",
}


BUSINESS_RULES = [
    ("BR-WS-001", "Workspace module hiện chỉ có Enterprise Workspace; không tồn tại non-enterprise workspace type hoặc workspace-type branching."),
    ("BR-WS-002", "Enterprise Workspace luôn có ít nhất một Owner và có thể có Owner/Admin/Member với MembershipType Internal hoặc External."),
    ("BR-WS-003", "Owner có quyền quản trị cao nhất; Admin quản trị vận hành nhưng không quản lý Owner, billing hoặc delete workspace."),
    ("BR-WS-004", "Admin không được quản lý Owner, không đổi role của Admin khác và không promote Member lên Admin theo logic service hiện tại."),
    ("BR-WS-005", "Invitation cho cùng email đang pending khi resend phải làm token cũ thành replaced và cấp token mới."),
    ("BR-WS-006", "Verified domain phải so khớp theo exact domain equality; subdomain chỉ hợp lệ khi cấu hình cho phép."),
    ("BR-WS-007", "External Member mặc định không thấy directory nội bộ, settings, document, transcript hoặc artifact toàn workspace."),
    ("BR-WS-008", "External meeting exception chỉ cho truy cập tài nguyên của meeting họ tham gia trong grace period cấu hình."),
    ("BR-WS-009", "Document ACL áp dụng deny-overrides và default deny cho sensitive/external/pending ingestion."),
    ("BR-WS-010", "Document Owner hoặc Workspace Owner/Admin mới được chỉnh policy, metadata nhạy cảm và trạng thái approval."),
    ("BR-WS-011", "Không hard-delete workspace; workspace inactive/soft-delete vẫn giữ khóa ngoại và audit trail."),
    ("BR-WS-012", "Không duplicate AI chunks trong workspace schema; vector documents/chunks thuộc AI schema/service."),
    ("BR-WS-013", "Chỉ người dùng có vai trò Owner hoặc Admin mới có quyền truy cập Dashboard thống kê và nhật ký hoạt động của Workspace."),
    ("BR-WS-014", "Dữ liệu trên Dashboard phải được lọc tuyệt đối theo WorkspaceId hoạt động hiện tại để tránh rò rỉ dữ liệu chéo giữa các tenant."),
    ("BR-WS-015", "Các số liệu về phòng họp dịch thuật phải được truy vấn từ TranslationRoom Service qua gRPC, không truy vấn trực tiếp cơ sở dữ liệu."),
    ("BR-WS-016", "Nhật ký hoạt động (Activity Logs) trên Dashboard chỉ ghi nhận và hiển thị các hành động cấu hình hệ thống, thay đổi thành viên và tài liệu nhạy cảm."),
    ("BR-WS-017", "Quyền tạo meeting là thuộc tính của membership; hệ thống dùng `workspace_members.can_create_meetings` cho per-member override, không dùng danh sách userId allow/deny trong `workspace.settings` JSONB."),
    ("BR-WS-018", "Một user chỉ được xem là active app user khi có ít nhất một active membership trong một Enterprise Workspace đang active."),
    ("BR-WS-019", "Khi workspace bị deactivated/soft-deleted, Workspace Service phải clear/invalidate active context liên quan và xác định user nào mất active workspace cuối cùng để báo Auth chuyển `SUSPENDED_NO_ACTIVE_WORKSPACE`."),
    ("BR-WS-020", "Remove/leave member luôn là soft-delete membership; sau mutation phải kiểm tra target user còn active workspace khác không. Nếu không còn, Workspace báo Auth suspend app account thay vì hard-delete hoặc admin-block."),
    ("BR-WS-021", "Accept invitation vào workspace active có thể re-activate account `SUSPENDED_NO_ACTIVE_WORKSPACE`, nhưng không được tự re-activate account `ADMIN_BLOCKED`, `DISABLED` hoặc soft-deleted."),
    ("BR-WS-022", "Mọi xóa workspace, member, document và account liên quan governance phải là soft-delete; không hard-delete dữ liệu phục vụ audit/history/legal trace."),
]

BUSINESS_RULES_SCOPE = (
    "Business rules là các ràng buộc quyết định nghiệp vụ mà service phải enforce khi xử lý workspace. "
    "Phần này không trộn với Non-functional Requirements: nếu một dòng mô tả security/performance/availability chung thì nằm ở NFR; "
    "nếu dòng đó quyết định ai được làm gì, trạng thái nào hợp lệ, dữ liệu nào bị chặn hoặc quan hệ nào bắt buộc thì nằm ở Business Rules. "
    "Danh sách dưới đây được tổng hợp từ specs và đối chiếu với code WorkspaceService, WorkspaceInvitationService, WorkspaceMemberService, WorkspaceDocumentService và DocumentAccessEvaluator."
)


BUSINESS_RULE_SOURCE_DETAILS = {
    "BR-WS-001": "WT-139/AC + code: không có workspace type branch trong Workspace domain/model",
    "BR-WS-002": "WT-141/157 + code: WorkspaceMember role/membership type, owner guard",
    "BR-WS-003": "WT-141 + code: role extension IsOwner/IsAdmin/IsOwnerOrAdmin",
    "BR-WS-004": "WT-141 + code: WorkspaceMemberService.ChangeMemberRoleAsync admin restrictions",
    "BR-WS-005": "WT-140 + code: WorkspaceInvitationService resend/replaced token logic",
    "BR-WS-006": "WT-157 + code: EmailAddress/domain matching, verified domain settings",
    "BR-WS-007": "WT-157 + code: external member guard in member/settings/document access flows",
    "BR-WS-008": "WT-159 + spec: meeting artifact participant/grace-period policy",
    "BR-WS-009": "WT-158 + code: DocumentAccessEvaluator deny-overrides/default deny logic",
    "BR-WS-010": "WT-158 + code: WorkspaceDocumentService metadata/policy/approval checks",
    "BR-WS-011": "WT-139/141 + code/schema: IsActive/DeletedAt/DeletedBy soft-delete pattern",
    "BR-WS-012": "DB/AI Guardrails + infrastructure schema: chunks/vector belong to AI service/schema",
    "BR-WS-013": "Workspace role requirements; code: WorkspaceDashboardController RBAC guards",
    "BR-WS-014": "Tenant isolation policy; code: IWorkspaceContext.WorkspaceId enforcement",
    "BR-WS-015": "Microservice database boundary; code: TranslationRoomGrpcClient integration",
    "BR-WS-016": "Dashboard audit policy; code: WorkspaceDashboardService log filtering",
    "BR-WS-017": "WT-159 design decision: quyền tạo meeting theo từng member cần query/audit/validate trực tiếp trên workspace_members",
    "BR-WS-018": "Enterprise-only account principle; Workspace active membership is the source of truth for app eligibility",
    "BR-WS-019": "Workspace deactivation lifecycle; Auth account status sync requirement",
    "BR-WS-020": "Workspace member soft-delete lifecycle; Auth suspension sync requirement",
    "BR-WS-021": "Invitation lifecycle + Auth account status taxonomy",
    "BR-WS-022": "Soft-delete-only governance policy across Workspace/Auth audit boundaries",
}


NFRS = [
    ("NFR-WS-001", "Security", "Tất cả endpoint yêu cầu JWT hợp lệ, trừ invitation preview an toàn không lộ token hash."),
    ("NFR-WS-002", "Security", "Downstream services chỉ tin workspace context nội bộ đã ký, không tin header do client tự gửi."),
    ("NFR-WS-003", "Privacy", "Không lộ dữ liệu workspace khác; mọi query phải scope theo workspace_id và active membership."),
    ("NFR-WS-004", "Performance", "List workspace/member/document dùng phân trang; DB query mục tiêu dưới 50ms cho list cốt lõi."),
    ("NFR-WS-005", "Availability", "Redis/RabbitMQ/AI ingestion failure không làm mất metadata upload; lỗi worker phải retry/audit/dead-letter được."),
    ("NFR-WS-006", "Compliance", "Sensitive document view/download/delete phải ghi audit action, actor, IP, user agent, metadata."),
    ("NFR-WS-007", "Maintainability", "Workspace Service không cross-join database của Auth/TranslationRoom; dùng gRPC/client boundary."),
    ("NFR-WS-008", "Scalability", "Document ingestion chạy bất đồng bộ qua Redis Stream kết hợp RabbitMQ với prefetch/concurrency limit, retry và dead-letter để tránh nghẽn CPU/AI."),
    ("NFR-WS-009", "Integrity", "Mọi workspace tạo mới phải có Owner; mọi document phải có workspace_id và storage metadata hợp lệ."),
    ("NFR-WS-010", "Cryptography", "Local encrypted file phải verify HMAC trước khi decrypt và dùng constant-time compare."),
    ("NFR-WS-011", "Frontend security", "Workspace UI phải kế thừa security header pattern từ web: X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy strict-origin-when-cross-origin."),
    ("NFR-WS-012", "Frontend performance", "Workspace UI phải dùng request timeout, loading skeleton, pagination và cache immutable cho static assets; không block thao tác chính vì panel phụ tải chậm."),
    ("NFR-WS-013", "Frontend resilience", "Workspace UI phải preserve form input sau network error, dùng retry rõ ràng, refresh-token queue để tránh nhiều request refresh đồng thời và redirect login khi session hết hạn."),
    ("NFR-WS-014", "Frontend accessibility", "Workspace UI phải có label cho form fields, keyboard reachable controls, aria/status cho loading/error/success và không chỉ dựa vào màu để biểu thị trạng thái."),
    ("NFR-WS-015", "Testability", "Workspace backend requirement phải có automated test trace bằng xUnit/Microsoft.NET.Test.Sdk, NSubstitute cho service/controller isolation, Testcontainers PostgreSQL cho integration, coverlet.collector cho coverage và Postman/Newman-compatible collection cho API smoke/regression."),
    ("NFR-WS-016", "Regression control", "Mỗi thay đổi validation/constraint phải cập nhật unit test tương ứng và ít nhất một negative case cho API hoặc integration boundary."),
    ("NFR-WS-017", "Performance", "API lấy thống kê dashboard stats phải phản hồi trong thời gian dưới 100ms."),
    ("NFR-WS-018", "Performance", "Sử dụng Redis distributed cache cho dữ liệu gRPC và stats với TTL tối đa 5 phút để bảo vệ downstream services."),
    ("NFR-WS-019", "Security", "Kiểm tra quyền truy cập (RBAC) Dashboard tại API Gateway và ứng dụng backend dựa trên internal context."),
]

NON_FUNCTIONAL_SCOPE = (
    "Non-functional requirements chỉ mô tả thuộc tính chất lượng và ràng buộc vận hành: security, privacy, performance, availability, compliance, maintainability, scalability, integrity, cryptography và UI quality. "
    "Không đưa luồng nghiệp vụ, role permission hoặc state transition vào phần này; các nội dung đó nằm ở Functional Requirements hoặc Business Rules."
)


USE_CASES = [
    (
        "UC-01",
        "Tạo và chọn Workspace",
        "Authenticated User",
        "User có JWT hợp lệ.",
        "User gửi create workspace, hệ thống sinh slug, tạo workspace và owner membership trong transaction, sau đó user select workspace.",
        "Workspace active context được lưu vào Redis/session và downstream context được ký.",
        "Tên không hợp lệ, slug conflict không xử lý được, hoặc user select workspace không phải member thì bị từ chối.",
    ),
    (
        "UC-02",
        "Mời và chấp nhận thành viên",
        "Owner/Admin, Invited User",
        "Workspace là Enterprise Workspace và caller là Owner/Admin có quyền invite.",
        "Owner/Admin tạo invite, email nhận link, user preview invite, đăng nhập đúng email và accept.",
        "Invitation chuyển accepted, membership active được tạo/kích hoạt.",
        "Role Owner, email mismatch, token expired/revoked/replaced, internal domain không hợp lệ hoặc external disabled đều bị từ chối.",
    ),
    (
        "UC-03",
        "Quản lý thành viên và ownership",
        "Owner/Admin/Member",
        "Caller là active member trong workspace.",
        "Owner/Admin list member, đổi role, remove member; member/admin tự leave khi hợp lệ.",
        "Role/status được cập nhật đúng và audit lịch sử membership không mất.",
        "Admin quản lý Owner, owner cuối cùng rời/demote/remove, user ngoài workspace gọi API đều bị từ chối.",
    ),
    (
        "UC-04",
        "Cộng tác với External Member",
        "Owner/Admin, External Member",
        "Workspace có verified domain và policy external collaboration.",
        "Admin mời email ngoài domain, hệ thống ép role External Member; external tham gia meeting được chỉ định.",
        "External chỉ xem contact Owner/Admin và tài nguyên meeting họ tham gia trong grace period.",
        "External bị chặn settings, directory nội bộ, document/transcript/artifact ngoài scope meeting.",
    ),
    (
        "UC-05",
        "Upload, phân loại và truy cập document",
        "Owner/Admin/Member, Worker",
        "User là active member và file đúng policy.",
        "User upload document, Workspace lưu metadata, publish Redis event, worker ingestion phân loại AI/sensitive, ACL quyết định truy cập.",
        "Document có audit trail, status/ingestion_status đúng và AI chỉ dùng document eligible/approved.",
        "File sai định dạng/quá size, document pending với member thường, explicit deny hoặc sensitive default deny đều bị chặn.",
    ),
    (
        "UC-06",
        "Governance cho meeting và artifact",
        "Host, Workspace Service, TranslationRoom Service",
        "Host có active workspace và quyền tạo meeting.",
        "TranslationRoom gọi Workspace gRPC để validate member, allowed languages, max active rooms; artifact nhận retention policy.",
        "Meeting/artifact đúng workspace boundary và retention/audit policy.",
        "Host external không được tạo internal meeting, language ngoài policy hoặc workspace inactive thì bị từ chối.",
    ),
]


FUNCTIONAL_TEST_CASES = [
    ("FR-WS-001", "Tạo Enterprise Workspace với name hợp lệ, domain công ty hợp lệ, Auth trả user/Owner role; response có workspace id/slug và member Owner.", "Tên có dấu/khoảng trắng sinh slug ổn định; slug trùng được ResolveSlugCollision; RequireVerifiedDomainForInternal=true nhưng request không truyền domain thì dùng domain email người tạo.", "Name rỗng, user không tồn tại, email user sai format, public domain, domain đã verify ở workspace khác, user đã là internal ở enterprise khác, thiếu Owner role."),
    ("FR-WS-002", "Đăng ký/login user không tự phát sinh workspace cá nhân; chỉ khi gọi create workspace mới có workspace.", "User thuộc nhiều workspace chỉ thấy danh sách workspace mình là active member; active workspace chưa chọn thì downstream yêu cầu chọn workspace.", "Không được xuất hiện endpoint/DB field/flow tạo personal workspace mặc định hoặc workspace type ngoài Enterprise."),
    ("FR-WS-003", "List workspace trả đúng page/pageSize/search và chỉ gồm workspace user còn active member.", "Search không dấu/có dấu, page vượt total trả empty page hợp lệ, member removed không còn thấy workspace.", "User chưa auth, repository lỗi, query page/pageSize không hợp lệ, workspace soft-delete/inactive không được lộ."),
    ("FR-WS-004", "User chọn workspace mình là active member; cache lưu workspace id, role, membership type.", "Role đổi sau lần select cần cache được refresh khi select lại; membership type xác định theo verified domain/settings.", "User không phải member, member đã removed, workspace không tồn tại, Redis/cache lỗi phải trả lỗi rõ hoặc không làm sai context."),
    ("FR-WS-005", "Gateway/Auth truyền internal workspace context đã ký cho downstream; service downstream đọc UserId/WorkspaceId/Role.", "Context hết hạn hoặc role vừa đổi phải bị refresh theo session policy.", "Client tự spoof header workspace context, chữ ký sai, workspace mismatch hoặc missing active workspace đều bị từ chối."),
    ("FR-WS-006", "Owner/Admin mời internal/external theo policy; pending invitation có token raw chỉ dùng để gửi email trực tiếp cho receiver, token_hash lưu DB.", "Resend cùng email làm pending cũ thành REPLACED; email mới được gửi trực tiếp tới receiver; language email fallback theo preferred language rồi workspace default rồi en.", "Member mời user, Admin assign Owner, external disabled, external non-Member role, internal domain chưa verified, membershipType sai, Email service fail thì invite vẫn pending và cần retry/delivery warning."),
    ("FR-WS-007", "Pending invite accept thành ACCEPTED; revoke thành REVOKED; resend thành REPLACED; expired token thành EXPIRED khi kiểm tra.", "Preview pending nhưng ExpiresAt đã qua trả currentStatus EXPIRED; revoked/replaced giữ audit trạng thái.", "Accept token ACCEPTED/REVOKED/EXPIRED/REPLACED bị reject với InvalidState."),
    ("FR-WS-008", "Authenticated user có email trùng invited email accept thành công và tạo workspace member.", "Email so khớp case-insensitive; user đã có preferred language không ảnh hưởng accept.", "Email mismatch, token rỗng, token hash không tìm thấy, user email claim thiếu, đã là member, internal member đã thuộc enterprise khác khi policy yêu cầu."),
    ("FR-WS-009", "Internal Owner/Admin list members thấy email, role, status; Internal Member thấy danh sách active phù hợp policy.", "Search theo tên/email, role name cache theo roleId, page vượt total trả empty.", "External Member gọi list directory bị Forbidden; user không active member; workspace không tồn tại."),
    ("FR-WS-010", "Owner/Admin remove Member soft-delete: RemovedAt/RemovedBy/Status=Removed; removed member mất quyền truy cập.", "Member/Admin tự leave thành công; remove user đã removed trả not found.", "Member remove người khác, Admin remove Owner, target không tồn tại, repository save lỗi."),
    ("FR-WS-011", "Workspace còn nhiều Owner thì một Owner có thể demote/leave theo rule; transfer ownership đổi OwnerId và role.", "Current owner sau transfer bị chuyển Admin; target active non-external member thành Owner.", "Last Owner leave/demote bị reject; transfer bởi non-owner; target external/removed/non-member; thiếu role Owner/Admin trong Auth."),
    ("FR-WS-012", "Internal invite/create enforces one Internal Home Workspace when verified-domain policy is active.", "External membership in another workspace does not block the internal-home rule when policy allows it.", "User already Internal in another domain-verified Enterprise Workspace is rejected for a second Internal membership."),
    ("FR-WS-013", "External email được mời khi AllowExternalCollaboration=true và role Member.", "Workspace bật AllowSubdomains chỉ ảnh hưởng internal domain matching nếu được cấu hình; external vẫn bị hạn chế role.", "External disabled, external role Admin/Owner, internal domain giả mạo, public domain verify."),
    ("FR-WS-014", "External Member chỉ truy cập tài nguyên meeting/document được grant trực tiếp và không quản trị workspace.", "External tham gia nhiều enterprise workspace vẫn giữ membership boundary riêng.", "External xem directory/settings/policies/toàn bộ artifacts bị reject."),
    ("FR-WS-015", "Upload document lưu metadata, storage key, owner/uploader, status/ingestion status, audit upload.", "Owner/Admin upload active+pending ingestion; Member upload pending approval+awaiting approval.", "Workspace inactive/deleted, user không phải member, file metadata thiếu, save DB lỗi."),
    ("FR-WS-016", "Evaluate ACL với explicit allow cho view/download trả success khi không có deny.", "Nhiều policy cùng lúc: deny thắng allow; owner/admin override theo rule quản trị nếu policy cho phép.", "Explicit deny, sensitive default deny, pending ingestion với member thường, policy subject mismatch."),
    ("FR-WS-017", "Owner/Admin approve pending document: status active, ingestion pending, AiEligible true, publish event.", "Reject pending document: status rejected, AiEligible false, audit reject.", "Approve document không pending, non-owner/admin approve, document không thuộc workspace, deleted document."),
    ("FR-WS-018", "Delete/archive document set DeletedAt/DeletedBy/AiEligible=false và publish invalidation event.", "Download/list sau delete không trả document; audit vẫn giữ.", "AI/RAG dùng document deleted/archived, delete bởi non-owner/non-doc-owner, delete document không tồn tại."),
    ("FR-WS-019", "Upload accepted extension/size theo policy và tạo storageKey đúng workspace/document.", "Tên file dài/ký tự đặc biệt vẫn lưu metadata an toàn; extension normalize lower-case.", "Unsupported extension, quá size, missing fileName/storage metadata, content type giả mạo."),
    ("FR-WS-020", "Local storage encrypt-then-MAC, verify HMAC trước decrypt, trả plaintext chỉ khi MAC hợp lệ.", "Rotate key theo workspace cần đọc được version cũ nếu có metadata version.", "HMAC sai, key thiếu, ciphertext corrupt, timing leak khi compare MAC."),
    ("FR-WS-021", "DocumentUploaded event được publish qua Redis Stream và chuyển tiếp/đồng bộ với RabbitMQ cho worker AI; worker cập nhật ingestion status.", "Redis tạm lỗi không rollback metadata; RabbitMQ retry/dead-letter giữ event idempotent.", "Worker fail scan thì document fail-safe: IsSensitive=true, AiEligible=false, ingestion failed."),
    ("FR-WS-022", "TranslationRoom gọi Workspace boundary để validate member/policy trước create/join room.", "Workspace policy thay đổi giữa lúc room setup cần validate lại trước start.", "TranslationRoom cross-join DB workspace, missing workspace context, external host tạo internal meeting."),
    ("FR-WS-023", "Workspace policy chặn max active rooms/target language ngoài allow-list và áp retention cho artifacts.", "Room scheduled chuyển live gần giờ vẫn check policy hiện tại.", "Inactive workspace, language không allowed, max active rooms exceeded, retention missing cho artifact sensitive."),
    ("FR-WS-024", "Workspace glossary theo business_domain/source/target/term ảnh hưởng AI/translation context.", "Duplicate term khác target_language hợp lệ; inactive term không đưa vào prompt.", "Duplicate cùng workspace/domain/source/target/term, unsupported language, user thiếu quyền quản lý glossary."),
    ("FR-WS-027", "Owner/Admin bật `can_create_meetings=true` cho một Internal Member; TranslationRoom validate thành công và cho tạo room nếu các policy khác pass.", "Admin/Owner tắt quyền của một Member nhưng không đổi role; member vẫn xem tài nguyên được phép nhưng không tạo meeting.", "External default false, removed member, non-owner/admin toggle, hoặc member có `can_create_meetings=false` tạo room đều bị 403."),
]


FUNCTIONAL_LAYER_MAPPING = [
    ("FR-WS-001", "POST /api/v1/workspaces nhận CreateWorkspaceRequest, lấy UserId từ JWT, trả ApiErrorResponse khi fail.", "WorkspaceService.CreateWorkspaceAsync validate name/user/email/domain, sinh slug, resolve Owner role, tạo Workspace+Owner member trong transaction.", "Workspace, WorkspaceMember, WorkspaceVerifiedDomain, EmailAddress, WorkspaceConfiguration, WorkspaceMemberRole.Owner.", "WorkspaceRepository, WorkspaceMemberRepository, generic repository verified_domains, UnitOfWork, AuthIdentityGrpcClient, PostgreSQL workspace schema."),
    ("FR-WS-002", "Không expose route personal/default workspace; routes chỉ tạo/list/select Enterprise Workspace.", "Service không gọi auto-provision personal workspace; logic chỉ chạy khi explicit create/select.", "Không có WorkspaceType enum/column; Enterprise behavior qua Workspace settings và MembershipType.", "DB không có workspace_type; migration/schema chỉ lưu workspaces, members, invitations, verified_domains."),
    ("FR-WS-003", "GET /api/v1/workspaces nhận GetWorkspacesQuery page/pageSize/search.", "WorkspaceService.GetWorkspacesAsync gọi repository theo user active membership, map role từ Auth.", "Workspace active membership, WorkspaceMember.RemovedAt, role vocabulary.", "WorkspaceRepository.GetWorkspacesForUserAsync, EF Core query, pagination trên PostgreSQL."),
    ("FR-WS-004", "POST /api/v1/workspaces/{id}/select.", "SelectWorkspaceAsync xác thực active member, resolve role/membership type, set active workspace cache.", "MembershipType Internal/External, WorkspaceConfiguration verified-domain rules.", "WorkspaceCacheService dùng Redis; repository đọc workspace/member; AuthIdentityGrpcClient resolve user/role."),
    ("FR-WS-005", "API/Gateway nhận JWT và downstream dùng internal signed context, không nhận client spoof context.", "Middleware/context service chuẩn hóa UserId/ActiveWorkspaceId/Role cho downstream calls.", "IWorkspaceContext contract, role/membership vocabulary.", "Redis/session cache active workspace, shared middleware, signing/verifier config."),
    ("FR-WS-006", "POST /api/v1/workspaces/{workspaceId}/invitations.", "WorkspaceInvitationService.InviteMemberAsync validate inviter role, email, membershipType, domain/external policy, token hash.", "WorkspaceInvitation, InvitationStatus.PENDING/REPLACED, MembershipType, EmailAddress.", "WorkspaceInvitationRepository, WorkspaceMemberRepository, verified_domain repository, AuthIdentityGrpcClient, UnitOfWork."),
    ("FR-WS-007", "GET/DELETE invitation endpoints và accept/preview endpoints trả trạng thái phù hợp.", "List/Revoke/Preview/Accept methods chuyển trạng thái PENDING/ACCEPTED/REVOKED/EXPIRED/REPLACED.", "InvitationStatus enum, WorkspaceInvitation.CheckAndHandleExpirationAsync.", "WorkspaceInvitationRepository by token hash/email/workspace, PostgreSQL persistence."),
    ("FR-WS-008", "POST /api/v1/workspaces/invitations/accept lấy UserId và Email claim.", "AcceptInvitationAsync hash token, validate status/expiry/email/domain/internal-elsewhere, tạo member và update invite.", "WorkspaceMember invitation member, MembershipType, EmailAddress, InvitationStatus.ACCEPTED.", "WorkspaceInvitationRepository.GetByTokenHashAsync, WorkspaceMemberRepository.AddAsync, UnitOfWork transaction."),
    ("FR-WS-009", "GET /api/v1/workspaces/{workspaceId}/members.", "WorkspaceMemberService.ListMembersAsync reject external caller, filter/search/page, mask email for non-admin members.", "WorkspaceMemberStatus, MembershipType.Internal, role extension IsOwnerOrAdmin.", "WorkspaceMemberRepository, WorkspaceRepository, AuthIdentityGrpcClient role/user lookup."),
    ("FR-WS-010", "DELETE /api/v1/workspaces/{workspaceId}/members/{userId}.", "RemoveMemberAsync handles self-leave, owner/admin removal, last-owner guard, soft-delete fields.", "WorkspaceMember.RemovedAt/RemovedBy/Status, WorkspaceMemberStatus.Removed.", "WorkspaceMemberRepository.CountActiveOwnersAsync, UnitOfWork SaveChanges."),
    ("FR-WS-011", "PUT role endpoint và POST transfer-ownership endpoint.", "ChangeMemberRoleAsync validates Admin/Member target roles and admin limits; TransferOwnershipAsync enforces owner-only and non-external target.", "Workspace.OwnerId, WorkspaceMember.RoleId, role extension Owner/Admin/Member, MembershipType.External.", "WorkspaceRepository.Update, WorkspaceMemberRepository, Auth role lookup via gRPC, UnitOfWork."),
    ("FR-WS-012", "Create/accept/invite APIs surface validation errors for Internal Home Workspace conflict.", "WorkspaceHelper.IsUserInternalMemberOfAnyEnterpriseWorkspaceAsync used during create/accept internal flows.", "MembershipType.Internal, verified-domain policy.", "WorkspaceMemberRepository/verified domain queries over PostgreSQL."),
    ("FR-WS-013", "Invitation/settings APIs expose AllowExternalCollaboration and membershipType External.", "InviteMemberAsync rejects external when disabled or role not Member; UpdateSettings enforces Owner-only external collaboration change.", "WorkspaceConfiguration.AllowExternalCollaboration, MembershipType.External.", "WorkspaceRepository settings JSONB, WorkspaceInvitationRepository."),
    ("FR-WS-014", "Member/document/settings endpoints gate external visibility.", "ListMembers rejects external; document access evaluator applies external boundary; settings update requires Owner/Admin.", "MembershipType.External, WorkspaceDocument ACL policy and role constants.", "Repositories + DocumentAccessEvaluator, audit tables for sensitive access."),
    ("FR-WS-015", "POST /api/v1/workspaces/{workspaceId}/documents.", "UploadDocumentAsync validates workspace/member, creates document state, saves metadata, audits, publishes event for owner/admin.", "WorkspaceDocument status/ingestion status, confidentiality flags, storage key convention.", "WorkspaceDocumentRepository, UnitOfWork.AuditAsync, RedisDocumentEventPublisher, URL provider."),
    ("FR-WS-016", "GET/download/policy endpoints call service/evaluator before returning document.", "DocumentAccessEvaluator.EvaluateAccessAsync computes permission with deny-overrides and defaults.", "WorkspaceDocumentAccessPolicy, WorkspaceDocumentPermissions, sensitive/pending status.", "WorkspaceDocumentAccessPolicyRepository, WorkspaceDocumentRepository."),
    ("FR-WS-017", "POST /documents/{documentId}/approve.", "ApproveDocumentAsync validates Owner/Admin and pending status, updates active/rejected state, publishes ingestion event when approved.", "WorkspaceDocumentStatus.pending_approval/active/rejected, IngestionStatus.pending.", "WorkspaceDocumentRepository, RedisDocumentEventPublisher, UnitOfWork audit."),
    ("FR-WS-018", "DELETE /documents/{documentId} and list/get/download filters.", "DeleteDocumentAsync soft-deletes, disables AI eligibility, publishes delete event, audits.", "WorkspaceDocument.DeletedAt/DeletedBy/AiEligible, audit action constants.", "WorkspaceDocumentRepository, Redis Stream + RabbitMQ bridge for invalidation, audit table."),
    ("FR-WS-019", "UploadDocumentRequest DTO carries file metadata; API rejects bad model state when validation is added.", "WorkspaceDocumentHelper.GenerateStorageKey and request mapper normalize document metadata.", "WorkspaceDocument fileName/fileExtension/storageKey/documentType fields.", "Storage provider boundary, PostgreSQL document metadata, future object storage adapter."),
    ("FR-WS-020", "Download endpoint returns signed/download URL or metadata after access check.", "Storage service must encrypt/decrypt local files around document download/upload boundary.", "Workspace-derived key metadata, confidentiality level.", "Local storage provider AES-256-CBC + HMAC-SHA512, key configuration, constant-time MAC compare."),
    ("FR-WS-021", "Upload/approve/delete APIs trigger publish operations but do not expose RabbitMQ internals.", "IWorkspaceDocumentEventPublisher abstraction publishes document events; background consumer processes ingestion asynchronously.", "Document event names DocumentUploaded/Deleted/Archived; ingestion status state machine.", "Redis Stream for local stream/cache and RabbitMQ for durable cross-service delivery, consumer group, retry/dead-letter."),
    ("FR-WS-022", "Workspace gRPC/client endpoints provide policy/member validation for TranslationRoom.", "Application client boundary validates workspace member, allowed languages, room limits.", "Workspace policy constants/settings, role/membership rules.", "TranslationRoomGrpcClient, AuthIdentityGrpcClient, no cross-service DB join."),
    ("FR-WS-023", "Workspace settings API exposes policy values consumed by room/artifact services.", "UpdateWorkspaceSettingsAsync validates Owner/Admin and Admin external-collaboration limitation.", "WorkspaceConfiguration max rooms/languages/retention/external collaboration.", "WorkspaceRepository JSONB settings, Redis active context/cache invalidation as needed."),
    ("FR-WS-024", "Terminology/glossary UI/API should operate under workspace manager permission.", "Application should validate duplicate terms, language pair, active/inactive lifecycle before AI use.", "WorkspaceKnowledgeGlossary unique key workspace_id+business_domain+source_language+target_language+term.", "workspace_knowledge_glossaries table, AI/translation prompt adapter boundary."),
    ("FR-WS-027", "PATCH /api/v1/workspaces/{workspaceId}/members/{userId}/meeting-permission nhận boolean canCreateMeetings; gRPC/member validation endpoint trả can_create_meetings cho TranslationRoom.", "WorkspaceMemberService validate caller Owner/Admin, target active member, external default false; TranslationRoom boundary reject khi false.", "WorkspaceMember.CanCreateMeetings là per-membership permission, độc lập với role; role vẫn quyết định quyền quản trị.", "Migration thêm workspace_members.can_create_meetings boolean NOT NULL; EF Core mapping; index tùy chọn (workspace_id, can_create_meetings) nếu cần list creator."),
]


FUNCTIONAL_IMPLEMENTATION_PLAN = [
    ("FR-WS-001", "Create workspace", "Controller nhận request/JWT; service validate user/domain/name; tạo workspace, slug, Owner membership trong UnitOfWork transaction; publish/cache invalidation sau commit; test rollback khi member insert fail."),
    ("FR-WS-002", "Enterprise-only workspace", "Không thêm WorkspaceType/Personal flow; giữ routes create/list/select Enterprise; audit code để không có auto-provision personal workspace; test account mới không tự có workspace."),
    ("FR-WS-003", "List workspaces", "Repository query theo active membership và workspace active; thêm search/page/pageSize validation; DTO trả role/membership type; test removed member và pagination edge."),
    ("FR-WS-004", "Select active workspace", "Validate membership active; resolve role từ Auth; ghi active context vào Redis; trả context DTO; refresh cache khi role/membership đổi; test stale membership/cache lỗi."),
    ("FR-WS-005", "Signed workspace context", "Middleware tạo/verify signed internal context gồm userId/workspaceId/role/membership; reject spoof/expired signature; downstream chỉ tin signed context, không tin client header."),
    ("FR-WS-006", "Invitation email delivery", "Owner/Admin gọi invite API; service validate role, email, domain/external policy; sinh raw token một lần, lưu token_hash; commit invitation; gọi Email/Notification client gửi email trực tiếp tới receiver với invite link; lưu trạng thái delivery/audit; nếu email fail thì invitation vẫn pending nhưng trả warning/retry state."),
    ("FR-WS-007", "Invitation lifecycle", "Model hóa Pending/Accepted/Revoked/Expired/Replaced; resend tạo token mới và mark pending cũ Replaced trong transaction; list/filter status; scheduled/preview expiry check; test revoke/accept race."),
    ("FR-WS-008", "Accept invitation", "Preview/accept hash raw token; validate pending, not expired, exact email match, internal enterprise rule; tạo hoặc reactivate member; mark invite Accepted; test email mismatch/expired/duplicate."),
    ("FR-WS-009", "List members", "Controller gọi service với query; service chặn External; repository list active hoặc manager view tùy role; map user profile/role từ Auth; mask fields nếu cần; test search/page/forbidden."),
    ("FR-WS-010", "Remove member", "Service validate caller Owner/Admin/self-leave; chặn Admin remove Owner và last Owner; set RemovedAt/RemovedBy/Status; invalidate member/workspace context; test soft-delete history."),
    ("FR-WS-011", "Ownership and role guard", "Implement ChangeRole/TransferOwnership transaction; resolve Owner/Admin/Member roleIds; chặn last Owner demote/leave; target transfer phải active internal; test missing role and external target."),
    ("FR-WS-012", "Internal home workspace constraint", "Helper query active internal memberships; enforce in create/accept internal invite; external membership does not block the internal-home rule when policy allows it; test mixed memberships."),
    ("FR-WS-013", "External collaboration", "Settings expose AllowExternalCollaboration; invite external only when enabled and role Member; Owner-only toggle external setting; test disabled/non-member role/public domain."),
    ("FR-WS-014", "External member boundary", "Centralize MembershipType.External guards in member/settings/document/artifact services; allow direct meeting-resource exception only; test directory/settings/document denial."),
    ("FR-WS-015", "Document library metadata", "Upload endpoint creates metadata/storage key/status; Owner/Admin active, Member pending approval; audit upload; publish event only after commit; test invalid workspace/member."),
    ("FR-WS-016", "Document ACL evaluator", "Implement evaluator order: document exists, member active, pending/sensitive checks, explicit DENY, explicit ALLOW, default rules, external meeting exception; unit test each branch."),
    ("FR-WS-017", "Document approval", "Approve/reject endpoint Owner/Admin only; pending-only transition; approve sets active + ingestion pending + publish event; reject records reason + aiEligible false; audit both."),
    ("FR-WS-018", "Archive/delete AI boundary", "Soft-delete/archive updates DeletedAt/DeletedBy/AiEligible false; publish invalidation to Redis/RabbitMQ; AI worker removes vector points; test event fail does not rollback metadata."),
    ("FR-WS-019", "Upload validation", "Add validation for extension, mime, size, filename, source type; normalize extension lowercase; reject oversized/unsupported; keep storage metadata consistent."),
    ("FR-WS-020", "Local encryption", "Storage provider encrypts AES-256-CBC and signs HMAC-SHA512; verify MAC before decrypt; store key/version metadata; test corrupt ciphertext/HMAC/key missing."),
    ("FR-WS-021", "AI ingestion messaging", "After DB commit publish workspace document events; Redis handles cache/local stream bridge, RabbitMQ durable delivery/retry/DLQ; worker idempotently updates ingestion/sensitivity."),
    ("FR-WS-022", "TranslationRoom boundary", "Expose Workspace validation client/gRPC response with membership, role, membershipType, can_create_meetings; TranslationRoom calls before create/join; no DB cross-join."),
    ("FR-WS-023", "Meeting policy settings", "Workspace settings DTO includes max active rooms, allowed languages, retention; service validates Owner/Admin update; room/artifact services consume through boundary/cache."),
    ("FR-WS-024", "Glossary", "Add CRUD/import/export service over workspace_knowledge_glossaries; enforce unique domain/source/target/term; expose active terms to AI/translation adapter."),
    ("FR-WS-025", "Dashboard stats", "Dashboard service aggregates workspace DB counts and TranslationRoom gRPC metrics; cache stats in Redis TTL <= 5 min; Owner/Admin only; degrade gracefully on gRPC fail."),
    ("FR-WS-026", "Dashboard activities", "Query audit/activity sources scoped by workspaceId; support search/filter/pagination; include settings/member/sensitive document actions; Owner/Admin only."),
    ("FR-WS-027", "Member meeting permission", "Migration adds can_create_meetings; mapper/DTO expose field; Owner/Admin toggle endpoint updates active member; TranslationRoom create-room check rejects false; tests cover internal/external/removed cases."),
]


BUSINESS_RULE_IMPLEMENTATION_PLAN = [
    ("BR-WS-001", "Enterprise-only", "Keep domain model without workspace_type; reject/avoid personal workspace routes; regression test no personal auto-provision."),
    ("BR-WS-002", "Owner/Admin/Member membership", "Use roleId from Auth and MembershipType field on workspace_members; validate active owner count in member mutations."),
    ("BR-WS-003", "Owner/Admin boundary", "Implement role extension checks in application services; keep Owner-only actions separate from Admin actions."),
    ("BR-WS-004", "Admin restrictions", "In ChangeMemberRole/RemoveMember reject Admin managing Owner/Admin or promoting Member to Admin; test each negative branch."),
    ("BR-WS-005", "Invitation resend", "When resend same email, mark previous pending invite Replaced, create new token_hash, send new email, old token invalid."),
    ("BR-WS-006", "Verified domain equality", "Normalize email/domain; exact match unless allow_subdomains true; public domains rejected; duplicate verified domain blocked by repository/index."),
    ("BR-WS-007", "External workspace visibility", "Apply External guard in members/settings/documents/artifacts routes and UI; return Forbidden with explicit reason."),
    ("BR-WS-008", "External meeting exception", "Document/artifact evaluator checks participant membership and grace period via TranslationRoom boundary before allowing direct resource access."),
    ("BR-WS-009", "Document deny-overrides", "Evaluator processes DENY before ALLOW and sensitive/default deny; policy tests assert DENY wins."),
    ("BR-WS-010", "Document policy mutation", "Only document owner or Owner/Admin can update metadata/policy/approval; audit each policy mutation."),
    ("BR-WS-011", "Soft-delete workspace", "Use is_active/deleted_at/deleted_by and filter active records; do not cascade-delete history/audit."),
    ("BR-WS-012", "AI chunks outside workspace DB", "Workspace stores metadata/source policy only; publish events to AI service for vector/chunk operations."),
    ("BR-WS-013", "Dashboard RBAC", "DashboardController/Service verifies Owner/Admin using internal context before stats/activities queries."),
    ("BR-WS-014", "Tenant isolation", "Every query includes workspaceId from active context/path; tests assert cross-workspace data is not returned."),
    ("BR-WS-015", "TranslationRoom metrics boundary", "Dashboard and meeting policy calls use TranslationRoom gRPC/client; no direct DB query to TranslationRoom schema."),
    ("BR-WS-016", "Activity log scope", "Activity service filters to settings/member/sensitive document actions; exclude noisy non-governance events."),
    ("BR-WS-017", "Meeting creator permission storage", "Persist per-member permission in workspace_members.can_create_meetings; do not implement settings JSONB userId allow/deny arrays; expose via DTO/gRPC and enforce at create-room."),
    ("BR-WS-018", "Active app user eligibility", "Workspace query/gRPC endpoint must be able to answer whether a user has at least one active membership in an active Enterprise Workspace; Auth uses this to decide ACTIVE vs SUSPENDED_NO_ACTIVE_WORKSPACE."),
    ("BR-WS-019", "Workspace deactivation sync", "On workspace deactivation/soft-delete, mark workspace inactive/deleted_at, clear Redis active contexts, publish WorkspaceDeactivated/UserWorkspaceEligibilityChanged events or call Auth gRPC for affected users."),
    ("BR-WS-020", "Member removal sync", "Remove/leave sets status/removed_at/removed_by, invalidates target active context, then checks remaining active memberships; if zero, request Auth set SUSPENDED_NO_ACTIVE_WORKSPACE."),
    ("BR-WS-021", "Invitation reactivation", "AcceptInvitation creates/reactivates membership only when workspace is active and invitation is valid; after success notify Auth to reactivate suspended account, but reject if Auth reports ADMIN_BLOCKED/DISABLED/SOFT_DELETED."),
    ("BR-WS-022", "Soft-delete only", "Keep deleted_at/deleted_by/status fields for workspace/member/document; never physically delete rows needed for FK, audit, meeting history, invitation history or legal trace."),
]


RABBITMQ_WORKFLOW = [
    ("1", "Publish domain event", "Workspace Application publishes `DocumentUploaded`, `DocumentDeleted`, `ArtifactCreated`, or `ArtifactRetentionExpired` through an application event publisher after DB metadata is committed.", "Use publisher confirms so the publisher knows RabbitMQ accepted the message; if confirm fails, record retry/audit without rolling back committed metadata."),
    ("2", "Route through exchange", "RabbitMQ receives the message at a durable topic exchange such as `workspace.events`.", "Use routing keys such as `workspace.document.uploaded`, `workspace.document.deleted`, `workspace.artifact.created`, `workspace.artifact.retention_expired`."),
    ("3", "Bind queues", "Queues bind to the exchange by routing key: `ai.document-ingestion`, `ai.embedding-invalidation`, `artifact.retention`, `audit.workspace-events`.", "Bindings decide which consumers receive each event; avoid direct service-to-service coupling."),
    ("4", "Consume with manual ack", "AI/artifact workers consume from queues with manual acknowledgements.", "Ack only after idempotency check, DB update and downstream side effects succeed. Nack/reject failed poison messages without requeue after retry limit."),
    ("5", "Control concurrency", "Consumers use prefetch/concurrency limit so document parsing, AI scanning and artifact cleanup do not overload CPU, storage or AI providers.", "Prefetch must be tuned separately for document ingestion and artifact cleanup."),
    ("6", "Retry transient failures", "Transient storage/AI/network failures are retried with bounded attempts.", "Use retry queues or delayed retry strategy; event payload must include eventId, documentId/artifactId, workspaceId and occurredAt for idempotency."),
    ("7", "Dead-letter permanent failures", "Messages rejected after retry limit, expired by TTL, or exceeding delivery limit are routed to a dead-letter exchange/queue.", "DLQ is monitored by ops; record failure reason and expose ingestion/artifact status as failed/requires_action."),
    ("8", "Reconcile state", "Scheduled reconciliation job compares DB records with queue/audit state to catch lost or stuck events.", "Metadata remains source of truth; RabbitMQ is delivery mechanism, not system of record."),
]


ARTIFACT_POST_MEETING_FLOW = [
    ("1. End room", "Host ends Translation Room; room status becomes ENDED and no new audio chunks should be accepted.", "TranslationRoom status and participant status.", "Reject late audio or route to ignored/degraded audit path."),
    ("2. Generate transcript", "Transcript service finalizes transcript segments and translations.", "transcript.transcripts, transcript_segments, transcript_translations.", "If transcript fails, artifact timeline shows failed and retry is available for Host/Admin."),
    ("3. Generate summary/report", "AI assistant creates summary, decisions, action items, risks and open questions using normalized transcript data.", "summary/report artifact metadata plus model metadata.", "If summary fails, transcript remains available; summary can be retried without reopening room."),
    ("4. Create artifact records", "System creates artifact records for transcript export, summary export and optional recording.", "translation_room.translation_room_artifacts or equivalent artifact table.", "Raw audio is not stored by default; only store when room policy explicitly enables recording."),
    ("5. Apply workspace retention", "RetentionUntil is calculated from Workspace ArtifactRetentionDays; raw audio uses shorter AudioRetentionDays.", "RetentionUntil, artifact type, workspace settings.", "Missing retention policy falls back to safe default and raises governance warning."),
    ("6. Apply access policy", "ArtifactAccess controls HostOnly, ParticipantsOnly or WorkspaceMembers; External Member access is limited to direct participant scope/grace period.", "artifact access metadata, room participants, workspace membership.", "Unauthorized users see locked/request-access state; sensitive access is audited."),
    ("7. Publish artifact events", "ArtifactCreated/ArtifactRetentionScheduled events are published via Redis/RabbitMQ for downstream indexing, notification and cleanup scheduling.", "RabbitMQ event payload, Redis cache invalidation if needed.", "Publish failure does not delete artifact metadata; reconciliation job retries."),
    ("8. User views artifacts", "Ended page shows generation timeline; Artifacts page shows Transcript, Summary, Action Items and Files tabs.", "artifact status, download URL, retention date.", "Not ready shows progress; failed shows retry; expired shows no longer available."),
    ("9. Retention cleanup", "Background job scans expired artifacts, deletes physical file from storage and soft-deletes/updates DB state while keeping audit metadata.", "storage object, artifact status, audit trail.", "Storage delete failure is retried and surfaced in ops queue; metadata is not silently removed."),
    ("10. Audit and traceability", "View/download/delete/retention actions write audit metadata with actor, workspaceId, artifactId, IP/user-agent when available.", "audit table/log events.", "Audit write failure must be logged and retried for sensitive artifacts."),
]


FUTURE_PROPOSED_SCOPE = [
    ("FP-WS-001", "Verified domain lifecycle", "Owner/Admin manages add, verify, enable, disable and remove company domains; unmanaged or revoked domains cannot grant internal access.", "Future/partially specified", "Linear WT-157 B2B Direction + spec 157"),
    ("FP-WS-002", "Domain verification edge cases", "Public domains, duplicate verified enterprise domains, disabled domains and existing non-matching members require explicit validation/migration behavior.", "Future/proposed hardening", "Linear WT-157 acceptance criteria"),
    ("FP-WS-003", "Document-to-AI boundary", "Workspace stores document metadata and AI eligibility while vector/chunk processing remains in AI service/schema; deleted/archived documents are excluded from retrieval.", "Future/proposed AI expansion", "Linear WT-158 B2B Direction + spec 158"),
    ("FP-WS-004", "Document approval and sensitive workflow", "Member uploads can require Owner/Admin approval before active ingestion; sensitive/default-deny/pending-ingestion states must be visible in API and UI.", "Future/proposed before broader rollout", "WT-158 approval addendum + AI guardrails"),
    ("FP-WS-005", "Native internal meeting governance", "Workspace governs who can create/join internal meetings, max active rooms, allowed languages and member-level CanCreateMeetings.", "Future/proposed before implementation", "Linear WT-159 + spec 159"),
    ("FP-WS-006", "Meeting document attachment", "Host/Admin can attach workspace documents to meetings only when document ACL and sensitive rules permit; participants receive time-bound access by meeting exception.", "Future/proposed before implementation", "WT-159 + WT-158"),
    ("FP-WS-007", "Post-meeting artifact lifecycle", "Transcript and summary artifacts are linked to workspace, receive RetentionUntil from ArtifactRetentionDays and are deleted by retention job when expired.", "Future/proposed before implementation", "Linear WT-159 + spec 159"),
    ("FP-WS-008", "Raw recording exclusion", "WT-159 scope does not create/store optional raw recording by default; any future recording feature needs separate consent, retention, audit and access rules.", "Future/proposed privacy guard", "Spec 159 updated scope"),
    ("FP-WS-009", "No cross-service DB joins", "TranslationRoom must validate workspace member/policy through Workspace gRPC/client boundary rather than querying workspace schema directly.", "Future/proposed architecture rule", "Linear WT-159 acceptance criteria"),
    ("FP-WS-010", "UI future surfaces", "Workspace UI must expose domain verification, document approval, meeting governance and artifact retention states before backend rollout to support implementation planning.", "Future/proposed UI spec", "Workspace UI spec + web .agents skills"),
]


BUSINESS_RULE_USER_STORY_TRACE = [
    ("WT-139", "BR-139-001..005", "Authenticated user creates/selects Enterprise Workspace; Owner membership and active context are explicit.", "Create/select workspace through real contracts; owner/membership bootstrap; downstream context contract."),
    ("WT-140", "BR-140-001..010", "Owner/Admin invites teammates/collaborators; invited user previews/accepts with exact email identity.", "Pending/accepted/revoked/expired/invalid/duplicate states and secure token storage are defined."),
    ("WT-141", "BR-141-001..008", "Owner/Admin manages members while preserving active Owner and soft-delete history.", "Permission denied, self-removal, owner protection and missing-member cases are defined."),
    ("WT-157", "BR-157-001..007", "Enterprise owner manages verified domains and separates Internal from External collaboration.", "Admin manages domains; unmanaged domains rejected; verification/revocation edge cases documented."),
    ("WT-158", "BR-158-001..008", "Workspace member uploads/accesses company knowledge safely under ACL, retention and AI boundary.", "Document CRUD/read contracts testable; permission, missing file, unsupported type and retention states handled."),
    ("WT-159", "BR-159-001..010", "Workspace governs native internal meetings and transcript/summary artifacts as future/proposed scope.", "Meetings organization-scoped; permissions for create/join/artifacts; artifacts linked to workspace; third-party optional."),
]


DB_ENTITIES = [
    ("workspace.workspaces", "Root tenant", "id PK, slug UK, owner_id FK->auth.users, settings JSONB, soft-delete fields"),
    ("workspace.workspace_members", "Membership assignment", "id PK, workspace_id FK, user_id FK->auth.users, role_id FK->auth.roles, can_create_meetings boolean, UNIQUE(workspace_id,user_id)"),
    ("workspace.workspace_invitations", "Invitation token lifecycle", "id PK, workspace_id FK, role_id FK->auth.roles, invited_by FK->auth.users, token_hash UK"),
    ("workspace.workspace_verified_domains", "Enterprise domain verification", "id PK, workspace_id FK, domain, status, verification_token, partial unique index on verified domain"),
    ("workspace.schema_migrations", "Workspace schema migration audit", "id PK, migration_key UK, checksum, status, started/completed timestamps"),
    ("workspace.workspace_documents", "Document library metadata", "id PK, workspace_id FK ON DELETE RESTRICT, storage fields, AI/retention/sensitivity fields"),
    ("workspace.workspace_document_access_policies", "Document ACL", "id PK, document_id FK ON DELETE CASCADE, workspace_id FK ON DELETE RESTRICT, subject/effect/permission"),
    ("workspace.workspace_document_audits", "Compliance audit trail", "id PK, document_id FK ON DELETE CASCADE, workspace_id FK ON DELETE RESTRICT, actor/action metadata"),
    ("workspace.workspace_knowledge_glossaries", "Workspace terminology", "id PK, workspace_id FK ON DELETE RESTRICT, UNIQUE(workspace_id,business_domain,source_language,target_language,term)"),
]

PHYSICAL_RELATIONSHIPS = [
    ("auth.users", "workspace.workspaces", "1", "0..N", "owner_id, created_by, updated_by, deleted_by"),
    ("workspace.workspaces", "workspace.workspace_members", "1", "0..N", "workspace_id"),
    ("auth.users", "workspace.workspace_members", "1", "0..N", "user_id, removed_by"),
    ("auth.roles", "workspace.workspace_members", "1", "0..N", "role_id"),
    ("workspace.workspaces", "workspace.workspace_invitations", "1", "0..N", "workspace_id"),
    ("auth.roles", "workspace.workspace_invitations", "1", "0..N", "role_id"),
    ("auth.users", "workspace.workspace_invitations", "1", "0..N", "invited_by"),
    ("workspace.workspaces", "workspace.workspace_verified_domains", "1", "0..N", "workspace_id"),
    ("auth.users", "workspace.workspace_verified_domains", "1", "0..N", "verified_by, created_by, updated_by"),
    ("workspace.workspaces", "workspace.workspace_documents", "1", "0..N", "workspace_id ON DELETE RESTRICT"),
    ("workspace.workspaces", "workspace.workspace_document_access_policies", "1", "0..N", "workspace_id ON DELETE RESTRICT"),
    ("workspace.workspace_documents", "workspace.workspace_document_access_policies", "1", "0..N", "document_id ON DELETE CASCADE"),
    ("workspace.workspaces", "workspace.workspace_document_audits", "1", "0..N", "workspace_id ON DELETE RESTRICT"),
    ("workspace.workspace_documents", "workspace.workspace_document_audits", "1", "0..N", "document_id ON DELETE CASCADE"),
    ("workspace.workspaces", "workspace.workspace_knowledge_glossaries", "1", "0..N", "workspace_id ON DELETE RESTRICT"),
]

MEETING_CREATOR_PERMISSION_DECISION = [
    ("Selected", "`workspace.workspace_members.can_create_meetings` boolean", "Quyền tạo meeting là thuộc tính của membership trong một workspace; dễ query, dễ audit, dễ trả về qua DTO/gRPC và không phình JSONB settings."),
    ("Rejected", "`workspace.settings.AllowedRoomCreatorUserIds` / `DisallowedRoomCreatorUserIds` JSONB arrays", "Không phù hợp cho per-member permission quy mô lớn: JSONB phình theo số member, khó truy vấn ngược, dễ stale userId khi remove/reinvite và phải deserialize settings khi check quyền."),
    ("Default", "Internal member true, External member false", "Owner/Admin/Member nội bộ được tạo meeting theo policy mặc định; External không được tạo meeting trừ khi có quyết định policy riêng sau này."),
    ("Enforcement", "TranslationRoom validates through Workspace boundary", "Create-room flow phải gọi Workspace API/gRPC để kiểm tra active membership, can_create_meetings, max active rooms và allowed languages; không cross-service DB join."),
    ("Migration", "Add non-null boolean column with backfill", "Thêm migration: column default false hoặc true có kiểm soát; backfill internal active members true, external false; removed members vẫn denied theo status/removed_at."),
]

DOC_CONTROL_ROWS = [
    ("Title", "Workspace Module Software Requirement Specification"),
    ("Version", "2.1"),
    ("Created by", "Ngô Xuân Hạnh Nhi"),
    ("Last updated", "2026-06-13"),
    ("Scope", "Workspace module only, cross-checked with backend, UI source-of-truth, selected web non-functional references, AI and infrastructure references."),
    ("Primary source", "warptalk-backend/specs + warptalk-infrastructure/scripts/init-db.sql + Workspace UI Google Doc source-of-truth."),
    ("Update rule", "Every material edit must update changelog, scope impact, source references and QA checklist."),
]

CHANGE_LOG = [
    ("1.0", "2026-06-11", "Codex", "Initial Workspace SRS and overview", "Consolidated workspace specs into MD/DOCX with diagrams."),
    ("1.1", "2026-06-11", "Codex", "ERD and document-control enhancement", "Regenerated ERD from infrastructure init-db.sql; added changelog, AI usage tracking, technology matrix, web route intent, limitations and QC checklist."),
    ("1.2", "2026-06-11", "Codex", "Enterprise-only correction", "Updated overview, business rules and use cases to match current Workspace Service code: Enterprise Workspace only, no non-enterprise workspace flows."),
    ("1.3", "2026-06-11", "Codex", "Functional testing and UI split", "Added author, per-functional happy/edge/unhappy cases, API/Application/Domain/Infrastructure mapping, Redis+RabbitMQ messaging, and separate Workspace UI spec."),
    ("1.4", "2026-06-11", "Codex", "RabbitMQ and artifact flow clarification", "Clarified Redis+RabbitMQ workflow using RabbitMQ official concepts and added WT-159 post-meeting artifact handling flow to module deliverables only."),
    ("1.5", "2026-06-11", "Codex", "Functional/Business/NFR separation", "Clarified section boundaries and added source traceability from workspace specs plus Workspace Service code for Functional Requirements and Business Rules."),
    ("1.6", "2026-06-12", "Codex", "Future/proposed scope and BR user stories", "Added business-rule user stories to workspace specs and captured future/proposed Workspace governance scope in SRS before implementation."),
    ("1.7", "2026-06-12", "Codex", "Testing and validation traceability", "Expanded happy/edge/unhappy testing details from workspace tests, validation/constraints and backend test toolchain."),
    ("1.8", "2026-06-12", "Antigravity", "Workspace Dashboard Spec & Purge Cleanup", "Added Workspace Dashboard specification and removed purged status from workspace lifecycle."),
    ("1.9", "2026-06-12", "Antigravity", "Detailed Functional Specs & Usecase Diagram Update", "Added detailed specifications for all FRs (FR-WS-001..026) and updated UML Use Case diagram."),
    ("2.0", "2026-06-12", "Codex", "Meeting creator permission data decision", "Added decision to persist per-member meeting creation permission as workspace_members.can_create_meetings instead of Workspace settings JSONB allow/deny lists."),
    ("2.1", "2026-06-13", "Codex", "Enterprise-only account eligibility", "Added active workspace dependency rules, Auth suspension sync and soft-delete-only account/workspace lifecycle alignment."),
]

AI_USAGE_LOG = [
    ("2026-06-11", "Codex", "SRS generation", "Created Workspace overview/SRS from workspace specs and code inspection.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-11", "Codex", "SRS revision", "Added ERD standards research, physical ERD from init-db.sql, Google Doc-aligned control sections and technology matrix.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-11", "Codex", "Enterprise-only revision", "Corrected Workspace SRS/spec BR and use cases according to current code: single Enterprise Workspace model.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-11", "Codex", "Functional/UI revision", "Expanded Workspace backend functional testing matrix and separated UI screen specification from backend SRS.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-11", "Codex", "RabbitMQ/artifact revision", "Updated module deliverables to use RabbitMQ terminology and added WT-159 artifact post-meeting flow.", "Not available from available telemetry; record manually."),
    ("2026-06-11", "Codex", "Functional/BR/NFR revision", "Separated Functional, Business Rule and Non-functional scopes.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-12", "Codex", "Future/proposed revision", "Added future/proposed Workspace governance requirements.", "Not available from local API telemetry; record manually if platform reports usage."),
    ("2026-06-12", "Codex", "Testing revision", "Reviewed workspace/tests, WorkspaceDbContext constraints, validation services and backend test tooling.", "Not available from local API telemetry."),
    ("2026-06-12", "Antigravity", "Dashboard & Soft Delete Spec", "Integrated detailed dashboard specifications and updated the deactivation lifecycle.", "Not available from local API telemetry."),
    ("2026-06-12", "Antigravity", "Detailed FR Specifications & UML Use Case", "Added detailed specification sections for all FRs (FR-WS-001..026) and updated PIL and Mermaid use case diagrams.", "Not available from local API telemetry."),
    ("2026-06-12", "Codex", "Meeting creator permission decision", "Documented why can_create_meetings belongs to workspace_members and why JSONB settings allow/deny userId lists are not selected.", "Not available from local API telemetry."),
    ("2026-06-13", "Codex", "Enterprise-only account eligibility", "Added business rules for last active workspace loss, Auth `SUSPENDED_NO_ACTIVE_WORKSPACE` sync and invitation-based reactivation.", "Not available from local API telemetry."),
]

TECH_MATRIX = [
    ("Backend", "Language/runtime", "C#, .NET 10", "WorkspaceService API/Application/Domain/Infrastructure projects."),
    ("Backend", "API", "ASP.NET Core Controllers, JWT auth", "REST endpoints under /api/v1/workspaces."),
    ("Backend", "Inter-service", "gRPC", "Auth identity lookup, TranslationRoom policy/artifact integration."),
    ("Backend", "Persistence", "EF Core, Npgsql, PostgreSQL", "schema workspace; UUID v7; JSONB settings/policies."),
    ("Backend", "Messaging/cache", "Redis distributed cache + RabbitMQ", "Redis handles active workspace cache/local stream bridge; RabbitMQ handles durable document ingestion delivery, retry and dead-letter."),
    ("Backend", "Unit testing", "xUnit 2.9.3, Microsoft.NET.Test.Sdk 17.14.1", "Implemented in workspace/tests/WarpTalk.WorkspaceService.Tests for service, controller, middleware, ACL and ingestion behavior."),
    ("Backend", "Mocking", "NSubstitute 5.3.0", "Used by Workspace unit tests to isolate repositories, UnitOfWork, cache/event publisher, Auth client and URL provider."),
    ("Backend", "Integration testing", "Microsoft.AspNetCore.Mvc.Testing 10.0.0, Testcontainers.PostgreSql 4.0.0", "Used for Workspace invitation integration tests with real PostgreSQL container and WebApplicationFactory."),
    ("Backend", "Coverage", "coverlet.collector 6.0.4", "Collects coverage from dotnet test runs for Workspace test project."),
    ("Backend", "API smoke/E2E", "Postman collections, optional Newman runner", "Backend-level postman collections live under test/postman; Workspace module should add collection coverage for create/select/invite/member/document flows."),
    ("Web", "Framework", "Next.js App Router, React", "Workspace dashboard, terminology, billing, rooms/artifacts surfaces."),
    ("Web", "Data access", "Axios/TanStack Query pattern", "Workspace adapters should mirror backend response contracts."),
    ("Web", "Realtime", "SignalR client", "Room/meeting events; workspace pages consume downstream state."),
    ("Web", "UI/RBAC", "Role-aware routing required", "Owner/Admin/Member/External surfaces must be separated."),
    ("AI", "Ingestion", "Redis Stream workers + RabbitMQ consumers", "Document upload/archive/delete events; AI eligibility, sensitivity classification, retry and dead-letter handling."),
    ("AI", "PII/DLP", "Regex scanner now; Presidio target", "Production transition to Presidio NLP API with fallback scanner."),
    ("AI", "RAG/vector", "AI schema/vector store", "Workspace stores source metadata only; chunks/vector points live in AI domain."),
    ("AI", "Provider normalization", "Worker boundary normalization", "Backend/UI must not depend on provider-specific raw output."),
    ("Infrastructure", "Runtime", "Docker Compose", "Service orchestration for backend dependencies."),
    ("Infrastructure", "Database", "PostgreSQL, PgBouncer", "init-db.sql defines workspace physical schema and indexes."),
    ("Infrastructure", "Cache/messaging", "Redis + RabbitMQ", "Redis for cache/local streams/pub-sub/backplane; RabbitMQ for durable event delivery, retry and dead-letter."),
    ("Infrastructure", "Observability", "Prometheus, Grafana, Seq/OpenTelemetry collector", "Logs/traces/metrics should include workspace_id where relevant."),
    ("Infrastructure", "Backups", "PostgreSQL/Qdrant backup scripts", "Document metadata and vector data require coordinated backup policy."),
]

TEST_TOOLCHAIN = [
    ("xUnit 2.9.3", "workspace/tests/WarpTalk.WorkspaceService.Tests/*.cs", "Primary automated unit/integration test framework for Workspace backend."),
    ("Microsoft.NET.Test.Sdk 17.14.1", "Workspace test csproj", "Runs tests through dotnet test and Visual Studio test runner."),
    ("NSubstitute 5.3.0", "Workspace service/controller tests", "Mocks repositories, UnitOfWork, cache/event publisher and external service clients."),
    ("Microsoft.AspNetCore.Mvc.Testing 10.0.0", "Integration/BaseIntegrationTest.cs", "Hosts Workspace API in-memory through WebApplicationFactory for integration scenarios."),
    ("Testcontainers.PostgreSql 4.0.0", "Integration/BaseIntegrationTest.cs", "Starts isolated PostgreSQL for invitation integration tests and schema-backed verification."),
    ("coverlet.collector 6.0.4", "Workspace test csproj", "Coverage collection for CI/local dotnet test runs."),
    ("Postman 12.x collections", "test/postman", "Manual/E2E API smoke tests; current backend collections exist for auth/notification/translationroom/transcript and should be extended for Workspace."),
    ("Newman-compatible workflow", "planned CI usage", "Recommended CLI runner for Postman collections so Workspace API smoke tests can run in CI."),
]

VALIDATION_CONSTRAINTS = [
    ("Workspace creation", "WorkspaceService.CreateWorkspaceAsync; WorkspaceDbContext.workspaces", "Name is required; creator user/email must exist and be valid; public domains cannot be verified; verified domain cannot already belong to another workspace; slug is unique; workspace creates Owner membership atomically."),
    ("Active workspace selection", "WorkspaceService.SelectWorkspaceAsync; WorkspaceCacheService; InternalContextMiddleware", "Caller must be active member; selected context is stored in Redis/cache; signed internal context must have valid signature, non-expired token and non-blacklisted user."),
    ("Settings/domain update", "WorkspaceService.UpdateWorkspaceSettingsAsync; WorkspaceConfiguration", "Settings payload JSON must parse; Owner/Admin required; AllowExternalCollaboration mutation is Owner-only; public domains rejected; active verified-domain uniqueness must hold through workspace.workspace_verified_domains."),
    ("Invitation", "WorkspaceInvitationService; WorkspaceInvitationValidator; WorkspaceDbContext.workspace_invitations", "Email format valid; membershipType must be Internal/External; external collaboration must be enabled for External; External can only use Member role; internal requires verified domain when policy requires; token hash unique; pending-only revoke."),
    ("Invitation accept/preview", "WorkspaceInvitationService.AcceptInvitationAsync/PreviewInvitationAsync", "Token required; token hash must match; status must be pending; ExpiresAt not passed; authenticated email must match invited email; internal user cannot join a second domain-verified Enterprise Workspace as Internal when rule applies."),
    ("Members/ownership", "WorkspaceMemberService; WorkspaceDbContext.workspace_members", "Requester must be active member; External cannot list directory; Owner/Admin required for mutation; cannot remove/change Owner by Admin; cannot leave/demote last Owner; transfer target must be active internal member."),
    ("Meeting creator permission", "WorkspaceMemberService; WorkspaceDbContext.workspace_members", "Owner/Admin may toggle can_create_meetings per active member; Internal default true; External default false; removed members are denied regardless of flag; do not store per-user allow/deny lists in workspace.settings JSONB."),
    ("Document upload/approval", "WorkspaceDocumentService; WorkspaceDbContext.workspace_documents", "Caller must be active member; Owner/Admin upload active while Member upload pending approval; pending-only approval/rejection; document keeps workspace_id, storage key/provider, status and ingestion status."),
    ("Document ACL/access", "DocumentAccessEvaluator; WorkspaceDocumentAccessPolicy", "Document must exist; caller must be active workspace member; pending ingestion blocks non-owner/non-admin/non-doc-owner; deny overrides allow; sensitive document default deny; External requires meeting exception within grace period."),
    ("Document audit/events", "WorkspaceDocumentService; UnitOfWork.AuditAsync; RedisDocumentEventPublisher", "Upload/approve/reject/delete/policy changes audit action metadata; delete soft-deletes and publishes invalidation; event failure must not corrupt committed metadata."),
    ("AI ingestion", "DocumentAiIngestionConsumerService; WorkspaceConfiguration.AiUsagePolicy", "Document policy overrides workspace policy; workspace policy fallback applies; scanner failure is fail-safe: mark sensitive/not AI eligible and do not crash worker."),
    ("Database constraints", "WorkspaceDbContext", "workspace_members unique(workspace_id,user_id); invitations token_hash unique; workspace.workspace_verified_domains enforces active verified-domain uniqueness with a partial unique constraint; glossary unique(workspace_id,business_domain,source_language,target_language,term); document FKs restrict/cascade as modeled."),
]

EXISTING_TEST_COVERAGE = [
    ("WorkspaceServiceTests", "Create workspace", "Success bootstrap Owner; name empty; user already internal elsewhere; enterprise verified domains; duplicate domain; no domain; custom domains; public domain."),
    ("WorkspaceServiceTests", "List/select/detail/settings", "Paginated list; select saves cache; select non-member forbidden; get by id member/non-member/not found; settings parse/default; update settings Owner/Admin; non-manager forbidden; public domain rejected."),
    ("WorkspacesControllerTests", "Workspace API mapping", "201 create; 400 validation; 200 list/detail/select/settings; 403 for non-member and unauthorized settings update."),
    ("WorkspaceMemberServiceTests", "Member list/mutation", "Member list success; External list forbidden; Owner/Admin see removed/banned; non-member forbidden; owner removes member; admin cannot remove Owner; last Owner cannot leave; owner leaves when another Owner exists."),
    ("WorkspaceMemberServiceTests", "Role/ownership", "Owner promotes member; Admin cannot demote Owner; last Owner cannot demote self; non-owner cannot transfer; new owner not member/external rejected; valid transfer succeeds."),
    ("WorkspaceMemberServiceTests", "Meeting creator permission target", "Should add tests for Owner/Admin toggling can_create_meetings, non-manager forbidden, External default false, removed member denied and TranslationRoom validation returns false."),
    ("WorkspaceMembersControllerTests", "Member controller mapping", "Paginated list, remove member and change role controller success mapping."),
    ("WorkspaceInvitationServiceTests", "Invite lifecycle", "Invite success; resend replaces old pending; external disabled rejected; external non-Member role rejected; token not found; email mismatch; expired invite; internal already belongs elsewhere rejected; external joins multiple workspaces; valid accept succeeds; preview masks email; internal user can join another workspace as external; cannot join another as internal."),
    ("WorkspaceInvitationIntegrationTests", "Invitation integration", "Preview valid token with accountExists; accept valid user; internal enterprise conflict forbidden; workspace without verified domains succeeds."),
    ("WorkspaceDocumentServiceTests", "Documents", "Member upload pending approval; Admin upload active and publishes event; approve publishes; reject does not publish; download succeeds when access allowed; delete soft-deletes and publishes; get policies paginated; policy list access denied."),
    ("DocumentAccessEvaluatorTests", "ACL evaluation", "Document not found; non-member denied; pending ingestion blocks regular member; pending allows admin; deny overrides allow; allow policy grants; sensitive default deny; internal non-sensitive default allow; External requires meeting exception and grace period; policy management by role/document owner."),
    ("DocumentAiIngestionConsumerServiceTests", "AI ingestion", "PII marks sensitive/not eligible; DLP keyword marks sensitive/not eligible; document policy fallback to workspace settings; exception fail-safe does not crash."),
    ("WorkspaceConfigurationTests", "Workspace settings", "Safe defaults; normalize null/invalid JSON; retain valid values; serialize/deserialize AI policy with language-specific rules."),
    ("InternalContextMiddlewareTests", "Signed context", "Valid signed header binds context/user; invalid signature, expired token and blacklisted user return unauthorized; no header passes without context."),
    ("SlugGeneratorTests", "Slug helper", "ASCII, punctuation, whitespace, C#/.NET and Vietnamese diacritics normalized; collision appends suffix."),
]

PREDICTED_TEST_CASES = [
    ("TC-PRED-001", "Workspace create transaction rollback", "Repository/member add fails after workspace insert attempt.", "No workspace without Owner membership remains committed; transaction rolls back and returns controlled error."),
    ("TC-PRED-027", "Meeting creator permission toggle", "Owner/Admin toggles can_create_meetings for target active member.", "Field persists on workspace_members; DTO/gRPC returns updated value; TranslationRoom create-room check allows only when true."),
    ("TC-PRED-002", "Concurrent workspace creation same domain", "Two users create Enterprise Workspace with same verified domain concurrently.", "Only one verified domain reaches verified status; the other receives domain registered/unique constraint handling."),
    ("TC-PRED-003", "Slug collision high suffix", "Many existing slugs share same base.", "ResolveSlugCollision appends deterministic suffix and does not loop indefinitely."),
    ("TC-PRED-004", "Select workspace after membership removal", "User selected workspace, then Owner removes user.", "Next protected call/cache refresh rejects removed member and clears/invalidates active context."),
    ("TC-PRED-005", "Settings Owner-only external toggle", "Admin sends payload changing AllowExternalCollaboration with other valid fields.", "Request fails without partially applying forbidden toggle."),
    ("TC-PRED-006", "Invitation revoke race with accept", "Owner revokes pending invitation while invited user accepts token.", "Only one terminal state wins; no accepted membership from revoked token."),
    ("TC-PRED-007", "Invitation resend race", "Two resend requests for same email execute concurrently.", "Only latest pending token remains valid; older pending is replaced and token_hash uniqueness holds."),
    ("TC-PRED-008", "Email/domain normalization", "Invite/Create receives mixed-case email/domain and trailing spaces.", "Domain comparison is normalized; duplicates/public-domain checks still trigger."),
    ("TC-PRED-009", "Member role mutation stale role ids", "Auth role lookup misses Admin/Owner role or returns stale role id.", "Service returns validation error; no member role changes are persisted."),
    ("TC-PRED-010", "Owner transfer persistence", "Owner transfers ownership then old Owner tries Owner-only action.", "Old Owner no longer has Owner-only permission; new Owner can perform Owner-only action."),
    ("TC-PRED-011", "Document event publish failure", "Document metadata saved but Redis/RabbitMQ publish fails.", "Document state is not lost; ingestion status shows failed/requires_action or event retry is scheduled/audited."),
    ("TC-PRED-012", "Duplicate document policy", "Same subject/permission/effect is added twice.", "Service should reject duplicate or keep idempotent single effective policy; evaluator result remains deterministic."),
    ("TC-PRED-013", "Conflicting group/user policy", "MembershipType Internal allow but specific user deny.", "Deny-overrides returns AccessDeniedByPolicy."),
    ("TC-PRED-014", "External meeting exception boundary", "External participant accesses meeting document exactly at grace-period boundary.", "Define inclusive/exclusive boundary and assert consistent result; after boundary must deny."),
    ("TC-PRED-015", "Sensitive document owner path", "Document owner is regular Member and document is sensitive.", "Owner/document owner can manage/view only according to intended override; non-owner Member denied."),
    ("TC-PRED-016", "Retention expired document download", "Document retention_state expired but storage object still exists.", "Download denied and AI eligibility false even if physical file remains."),
    ("TC-PRED-017", "AI policy invalid JSON", "Document AiUsagePolicy contains invalid JSON.", "Worker logs warning and falls back to workspace policy or fail-safe default without crashing."),
    ("TC-PRED-018", "RabbitMQ idempotent redelivery", "Same DocumentUploaded event delivered twice.", "Worker processes idempotently; no duplicate audit/vector indexing side effects."),
    ("TC-PRED-019", "Postman Workspace smoke flow", "Run create workspace -> invite -> accept -> list members -> upload document -> approve -> download.", "Collection validates HTTP status, response code, workspace_id continuity and cleanup/negative cases."),
    ("TC-PRED-020", "gRPC boundary unavailable", "Auth/TranslationRoom gRPC dependency times out.", "Workspace returns controlled error, does not default to allow and logs dependency failure with correlation/workspace context."),
]

WEB_ROUTES = [
    ("/workspace/dashboard", "Workspace manager/owner", "Usage, members, rooms, governance overview."),
    ("/workspace/terminology", "Workspace manager/owner", "Glossary and terms by business domain."),
    ("/workspace/billing", "Workspace owner", "Plan, credits, usage and transactions."),
    ("/rooms", "Host/workspace", "Workspace-scoped room list."),
    ("/rooms/[id]/artifacts", "Host/workspace", "Transcript, summary and export artifacts under workspace policy."),
    ("/internal/dashboard", "Internal admin", "Tenants, platform health, AI operations; not a workspace member surface."),
    ("/internal/ai-ops", "Internal admin", "AI pipeline monitoring and operational review."),
]

LIMITATIONS = [
    ("High", "Workspace SRS now treats Workspace Service as present in backend code, but older system spec notes a prior mismatch where workspace APIs were implied by infrastructure. Keep this as a regression check when branches diverge."),
    ("High", "Full RBAC must be enforced consistently in backend and web middleware; web token-presence checks are not enough for Workspace surfaces."),
    ("High", "Redis stream contracts between Gateway/AI/Transcript must remain canonical before workspace document ingestion expands."),
    ("Medium", "Artifact retention and deletion workers need end-to-end verification against workspace policy."),
    ("Medium", "Response contracts across services should be standardized before web adapters depend on new workspace endpoints."),
    ("Medium", "Encoding issues in legacy specs/comments should be cleaned so Vietnamese requirements remain readable."),
    ("Low", "Status casing should be standardized across backend DTOs, database values and web adapters."),
]

QC_CHECKLIST = [
    ("Requirement", "Every new workspace behavior has FR/BR/NFR and at least one happy/unhappy case."),
    ("API", "Endpoint has request validation, typed success response and typed error response."),
    ("Security", "Auth, role and workspace membership policy are explicit; external member scope is tested."),
    ("Data", "Migration/DB changes include rollback note, indexes, FK/delete behavior and seed/default impact."),
    ("Performance", "List endpoints use bounded pagination and documented indexes."),
    ("AI", "Redis stream key/field names match canonical schema and have fallback/ retry behavior."),
    ("Web", "API adapters, route guards and TypeScript types are updated when backend contract changes."),
    ("Observability", "Logs/audits include user_id, workspace_id, document_id/room_id where available."),
    ("Docs", "Changelog, AI usage log and source traceability are updated in this SRS."),
]

DEFINITION_OF_DONE = [
    ("Backend", "Workspace API compiles, validates requests, enforces Owner/Admin/Member/External rules and has unit/integration tests."),
    ("Database", "workspace schema migration matches ERD, preserves audit history and avoids unsafe cascade deletes except document-owned ACL/audit rows."),
    ("Web", "Workspace dashboard, members, invitations, documents, terminology and billing surfaces map to real API contracts."),
    ("AI", "Document ingestion, sensitivity classification, AI eligibility and vector invalidation are observable and retry-safe."),
    ("Infrastructure", "PostgreSQL, Redis, storage and observability services are configured; backups cover metadata and vector dependencies."),
    ("Security", "Signed internal context, JWT auth, document ACL and external meeting exception are tested with negative cases."),
]


API_ROWS = [
    ("POST", "/api/v1/workspaces", "Tạo workspace và bootstrapping Owner"),
    ("GET", "/api/v1/workspaces", "List workspace theo pagination/search"),
    ("GET", "/api/v1/workspaces/{id}", "Xem chi tiết workspace nếu là member"),
    ("POST", "/api/v1/workspaces/{id}/select", "Chọn active workspace context"),
    ("GET/PUT", "/api/v1/workspaces/{id}/settings", "Xem/cập nhật workspace settings"),
    ("GET", "/api/v1/workspaces/{workspaceId}/members", "List active members"),
    ("DELETE", "/api/v1/workspaces/{workspaceId}/members/{userId}", "Remove/leave workspace bằng soft-delete"),
    ("PUT", "/api/v1/workspaces/{workspaceId}/members/{userId}/role", "Đổi role member"),
    ("PATCH", "/api/v1/workspaces/{workspaceId}/members/{userId}/meeting-permission", "Owner/Admin bật/tắt quyền tạo meeting per member bằng can_create_meetings"),
    ("POST", "/api/v1/workspaces/{workspaceId}/members/transfer-ownership", "Transfer ownership"),
    ("POST", "/api/v1/workspaces/{workspaceId}/invitations", "Tạo invite"),
    ("GET", "/api/v1/workspaces/{workspaceId}/invitations", "List invite"),
    ("DELETE", "/api/v1/workspaces/{workspaceId}/invitations/{invitationId}", "Revoke invite"),
    ("GET", "/api/v1/workspaces/invitations/preview", "Preview invite an toàn không cần JWT"),
    ("POST", "/api/v1/workspaces/invitations/accept", "Accept invite"),
    ("POST/GET", "/api/v1/workspaces/{workspaceId}/documents", "Upload/list documents"),
    ("GET/PATCH/DELETE", "/api/v1/workspaces/{workspaceId}/documents/{documentId}", "Xem/cập nhật metadata/xóa mềm document"),
    ("POST", "/api/v1/workspaces/{workspaceId}/documents/{documentId}/approve", "Approve/reject ingestion/sensitive decision"),
    ("GET", "/api/v1/workspaces/{workspaceId}/documents/{documentId}/download", "Download sau khi qua ACL"),
    ("POST/GET/DELETE", "/api/v1/workspaces/{workspaceId}/documents/{documentId}/policies", "Quản lý access policies"),
]


MERMAID = {
    "system": """flowchart LR
    Client["Web/Desktop Client"] --> Gateway["API Gateway/Auth"]
    Gateway -->|JWT + active workspace select| Workspace["Workspace Service"]
    Gateway -->|signed X-Internal-Context| Room["TranslationRoom Service"]
    Gateway --> Transcript["Transcript Service"]
    Workspace --> PG[("PostgreSQL schema: workspace")]
    Workspace --> Redis[("Redis cache + local streams")]
    Workspace --> Rabbit[("RabbitMQ exchanges + queues")]
    Workspace -->|gRPC identity lookup| Auth["Auth Service"]
    Workspace -->|gRPC room policy/artifact refs| Room
    Workspace --> Storage[("S3/MinIO/Local Storage")]
    Redis --> Rabbit
    Rabbit --> Worker["Document/Artifact AI Worker"]
    Worker --> AI["AI/Presidio/RAG services"]
    AI --> Vector[("AI vector schema / vector DB")]
""",
    "erd": """erDiagram
    AUTH_USERS ||--o{ WORKSPACES : owns_created_updates_deletes
    AUTH_USERS ||--o{ WORKSPACE_MEMBERS : member_user_removed_by
    AUTH_ROLES ||--o{ WORKSPACE_MEMBERS : assigned_role
    AUTH_ROLES ||--o{ WORKSPACE_INVITATIONS : invited_role
    AUTH_USERS ||--o{ WORKSPACE_INVITATIONS : invited_by
    AUTH_USERS ||--o{ WORKSPACE_VERIFIED_DOMAINS : verified_created_updated_by
    WORKSPACES ||--o{ WORKSPACE_MEMBERS : has
    WORKSPACES ||--o{ WORKSPACE_INVITATIONS : sends
    WORKSPACES ||--o{ WORKSPACE_VERIFIED_DOMAINS : verifies
    WORKSPACES ||--o{ WORKSPACE_DOCUMENTS : owns_restrict_delete
    WORKSPACES ||--o{ WORKSPACE_DOCUMENT_ACCESS_POLICIES : scopes_restrict_delete
    WORKSPACES ||--o{ WORKSPACE_DOCUMENT_AUDITS : scopes_restrict_delete
    WORKSPACES ||--o{ WORKSPACE_KNOWLEDGE_GLOSSARIES : owns_restrict_delete
    WORKSPACE_DOCUMENTS ||--o{ WORKSPACE_DOCUMENT_ACCESS_POLICIES : cascades_to
    WORKSPACE_DOCUMENTS ||--o{ WORKSPACE_DOCUMENT_AUDITS : cascades_to
    WORKSPACES {
      uuid id PK
      varchar name
      varchar slug UK
      uuid owner_id FK
      boolean allow_external_collaboration
      boolean require_verified_domain_for_internal
      boolean allow_subdomains
      jsonb settings
      boolean is_active
      timestamptz deleted_at
    }
    WORKSPACE_MEMBERS {
      uuid id PK
      uuid workspace_id FK
      uuid user_id FK
      uuid role_id FK
      varchar membership_type
      boolean can_create_meetings
      varchar status
      timestamptz removed_at
      unique workspace_user
    }
    WORKSPACE_INVITATIONS {
      uuid id PK
      uuid workspace_id FK
      varchar email
      uuid role_id FK
      varchar token_hash UK
      varchar status
      timestamptz expires_at
    }
    WORKSPACE_VERIFIED_DOMAINS {
      uuid id PK
      uuid workspace_id FK
      varchar domain
      varchar status
      varchar verification_token
      partial_unique verified_domain
    }
    WORKSPACE_DOCUMENTS {
      uuid id PK
      uuid workspace_id FK
      varchar storage_provider
      varchar storage_key
      varchar document_type
      boolean ai_eligible
      varchar ingestion_status
      boolean is_sensitive
      varchar retention_state
    }
    WORKSPACE_DOCUMENT_ACCESS_POLICIES {
      uuid id PK
      uuid document_id FK
      uuid workspace_id FK
      varchar subject_type
      uuid subject_id
      varchar permission
      varchar effect
    }
    WORKSPACE_DOCUMENT_AUDITS {
      uuid id PK
      uuid document_id FK
      uuid workspace_id FK
      uuid actor_id
      varchar action
      timestamptz action_at
    }
    WORKSPACE_KNOWLEDGE_GLOSSARIES {
      uuid id PK
      uuid workspace_id FK
      varchar business_domain
      varchar source_language
      varchar target_language
      varchar term
      unique glossary_term
    }
    AUTH_USERS {
      uuid id PK
      varchar email UK
    }
    AUTH_ROLES {
      uuid id PK
      varchar name UK
    }
""",
    "main": """flowchart TD
    A["Create/Select Workspace"] --> B["Invite or manage members"]
    B --> C{"Enterprise domain?"}
    C -->|internal| D["Assign Owner/Admin/Member"]
    C -->|external allowed| E["Force External Member"]
    C -->|external disabled| X["Reject invite"]
    D --> F["Create room / upload document"]
    E --> F
    F --> G{"Policy checks"}
    G -->|allowed| H["Persist metadata / meeting"]
    G -->|denied| Y["Return 403/400"]
    H --> I["Redis/gRPC downstream integration"]
    I --> J["Audit, retention, AI ingestion"]
""",
    "screen": """flowchart LR
    Login["Login"] --> Select["Workspace Switcher"]
    Select --> Dashboard["Workspace Dashboard"]
    Dashboard --> Members["Members"]
    Dashboard --> Invites["Invitations"]
    Dashboard --> Docs["Documents"]
    Dashboard --> Settings["Settings"]
    Docs --> DocDetail["Document Detail"]
    DocDetail --> Policies["Access Policies"]
    Dashboard --> Rooms["Meetings/Artifacts"]
    Settings --> Domains["Verified Domains"]
""",
    "invitation_state": """stateDiagram-v2
    [*] --> Pending: create invite
    Pending --> Accepted: accept with matching email
    Pending --> Revoked: owner/admin revoke
    Pending --> Expired: expires_at reached
    Pending --> Replaced: resend same email
    Accepted --> [*]
    Revoked --> [*]
    Expired --> [*]
    Replaced --> [*]
""",
    "document_state": """stateDiagram-v2
    [*] --> Uploaded
    Uploaded --> PendingIngestion: publish Redis event
    PendingInestionAlias --> PendingIngestion
    PendingIngestion --> AwaitingApproval: sensitive or policy requires review
    PendingIngestion --> Completed: classified safe
    AwaitingApproval --> Completed: owner/admin approves
    AwaitingApproval --> Rejected: owner/admin rejects
    Completed --> Archived: retention/archive action
    Completed --> Deleted: soft delete
    Archived --> Deleted: delete action
    Rejected --> Deleted
    Deleted --> [*]
""".replace("PendingInestionAlias --> PendingIngestion\n", ""),
    "workspace_account_state": """stateDiagram-v2
    [*] --> ActiveWorkspace: workspace created with owner
    ActiveWorkspace --> InactiveWorkspace: owner/admin deactivates
    ActiveWorkspace --> SoftDeletedWorkspace: delete workspace
    InactiveWorkspace --> ActiveWorkspace: restore/reactivate
    InactiveWorkspace --> SoftDeletedWorkspace: delete inactive workspace
    ActiveWorkspace --> MemberRemoved: remove/leave member
    MemberRemoved --> UserStillActive: user still has another active workspace
    MemberRemoved --> AuthSuspended: user has no active workspace left
    InactiveWorkspace --> AuthSuspended: affected user has no active workspace left
    SoftDeletedWorkspace --> AuthSuspended: affected user has no active workspace left
    AuthSuspended --> ActiveWorkspace: user accepts valid invitation into active workspace
    UserStillActive --> ActiveWorkspace
    SoftDeletedWorkspace --> [*]
""",
    "usecase": """flowchart TB
    Owner((Owner)) --> UC1["Create/select workspace"]
    Owner --> UC2["Manage settings/domains"]
    Owner --> UC3["Invite/manage members"]
    Owner --> UC4["Manage documents/ACL"]
    Admin((Admin)) --> UC3
    Admin --> UC4
    Member((Member)) --> UC5["Use rooms/documents by policy"]
    External((External Member)) --> UC6["Access invited meeting resources"]
    Worker((System Worker)) --> UC7["Ingest, classify, audit documents"]
    UC7 --> UC4
""",
}


def dedent(text: str) -> str:
    text = textwrap.dedent(text).strip()
    # Embedded Mermaid snippets contain zero-indented lines, so textwrap.dedent
    # cannot always remove the template margin. Strip one template indent layer.
    lines = [line[4:] if line.startswith("    ") else line for line in text.splitlines()]
    return "\n".join(lines) + "\n"


def functional_requirement_rows() -> list[tuple[str, str, str, str]]:
    rows = []
    for rid, area, desc, source in REQUIREMENTS:
        source_detail = FUNCTIONAL_SOURCE_DETAILS.get(rid, source)
        rows.append((rid, area, desc, source_detail))
    return rows


def business_rule_rows() -> list[tuple[str, str, str]]:
    rows = []
    for rid, desc in BUSINESS_RULES:
        source_detail = BUSINESS_RULE_SOURCE_DETAILS.get(rid, "Specs + Workspace code")
        rows.append((rid, desc, source_detail))
    return rows


def make_markdown() -> str:
    sources_table = "\n".join(
        f"| {sid} | [{path}]({path}) | {desc} |" for sid, path, desc in SOURCE_SPECS
    )
    doc_control_table = "\n".join(f"| {k} | {v} |" for k, v in DOC_CONTROL_ROWS)
    change_log_table = "\n".join(f"| {v} | {d} | {a} | {c} | {r} |" for v, d, a, c, r in CHANGE_LOG)
    ai_usage_table = "\n".join(f"| {d} | {a} | {s} | {w} | {u} |" for d, a, s, w, u in AI_USAGE_LOG)
    req_table = "\n".join(
        f"| {rid} | {area} | {desc} | {source} |" for rid, area, desc, source in functional_requirement_rows()
    )
    functional_implementation_plan_table = "\n".join(
        f"| {rid} | {function} | {plan} |"
        for rid, function, plan in FUNCTIONAL_IMPLEMENTATION_PLAN
    )
    functional_test_table = "\n".join(
        f"| {rid} | {happy} | {edge} | {unhappy} |" for rid, happy, edge, unhappy in FUNCTIONAL_TEST_CASES
    )
    layer_mapping_table = "\n".join(
        f"| {rid} | {api} | {application} | {domain} | {infrastructure} |"
        for rid, api, application, domain, infrastructure in FUNCTIONAL_LAYER_MAPPING
    )
    rabbitmq_table = "\n".join(
        f"| {step} | {activity} | {workspace_usage} | {rabbitmq_rule} |"
        for step, activity, workspace_usage, rabbitmq_rule in RABBITMQ_WORKFLOW
    )
    artifact_flow_table = "\n".join(
        f"| {step} | {activity} | {data} | {edge} |"
        for step, activity, data, edge in ARTIFACT_POST_MEETING_FLOW
    )
    future_scope_table = "\n".join(
        f"| {fid} | {capability} | {description} | {status} | {source} |"
        for fid, capability, description, status, source in FUTURE_PROPOSED_SCOPE
    )
    br_user_story_trace_table = "\n".join(
        f"| {ticket} | {rules} | {story} | {acceptance} |"
        for ticket, rules, story, acceptance in BUSINESS_RULE_USER_STORY_TRACE
    )
    br_table = "\n".join(f"| {rid} | {desc} | {source} |" for rid, desc, source in business_rule_rows())
    business_rule_implementation_plan_table = "\n".join(
        f"| {rid} | {rule_area} | {plan} |"
        for rid, rule_area, plan in BUSINESS_RULE_IMPLEMENTATION_PLAN
    )
    nfr_table = "\n".join(f"| {rid} | {area} | {desc} |" for rid, area, desc in NFRS)
    db_table = "\n".join(f"| `{name}` | {purpose} | {fields} |" for name, purpose, fields in DB_ENTITIES)
    rel_table = "\n".join(f"| `{p}` | `{c}` | {pc} | {cc} | {fk} |" for p, c, pc, cc, fk in PHYSICAL_RELATIONSHIPS)
    meeting_creator_decision_table = "\n".join(
        f"| {decision} | {option} | {rationale} |"
        for decision, option, rationale in MEETING_CREATOR_PERMISSION_DECISION
    )
    api_table = "\n".join(f"| {method} | `{route}` | {purpose} |" for method, route, purpose in API_ROWS)
    tech_table = "\n".join(f"| {area} | {topic} | {tech} | {usage} |" for area, topic, tech, usage in TECH_MATRIX)
    test_toolchain_table = "\n".join(
        f"| {tool} | {location} | {usage} |" for tool, location, usage in TEST_TOOLCHAIN
    )
    validation_constraints_table = "\n".join(
        f"| {area} | {source} | {constraint} |" for area, source, constraint in VALIDATION_CONSTRAINTS
    )
    existing_test_coverage_table = "\n".join(
        f"| {suite} | {scope} | {covered} |" for suite, scope, covered in EXISTING_TEST_COVERAGE
    )
    predicted_test_cases_table = "\n".join(
        f"| {tid} | {case} | {condition} | {expected} |" for tid, case, condition, expected in PREDICTED_TEST_CASES
    )
    web_route_table = "\n".join(f"| `{route}` | {audience} | {intent} |" for route, audience, intent in WEB_ROUTES)
    limitation_table = "\n".join(f"| {priority} | {item} |" for priority, item in LIMITATIONS)
    qc_table = "\n".join(f"| {area} | {check} |" for area, check in QC_CHECKLIST)
    dod_table = "\n".join(f"| {area} | {criterion} |" for area, criterion in DEFINITION_OF_DONE)
    uc_table = "\n".join(
        f"| {uid} | {name} | {actor} | {happy} | {unhappy} |"
        for uid, name, actor, pre, happy, post, unhappy in USE_CASES
    )

    return dedent(f"""
    # Workspace Module Requirements Overview

    **Ngôn ngữ:** Tiếng Việt  
    **Phạm vi:** WarpTalk Backend - Workspace Service và các tích hợp Auth, TranslationRoom, Transcript, AI/Document ingestion  
    **Ngày tạo:** 2026-06-11  
    **Deliverable song hành:** `workspace-software-requirement-specification.docx`

    ## 1. Document control

    | Field | Value |
    |---|---|
    {doc_control_table}

    ### 1.1 Change log

    | Version | Date | Author/AI | Change | Reason |
    |---|---|---|---|---|
    {change_log_table}

    ### 1.2 AI usage log

    | Date | AI/Actor | Scope | Work performed | Usage |
    |---|---|---|---|---|
    {ai_usage_table}

    ### 1.3 Rules for updating this file

    - Mọi thay đổi có ảnh hưởng requirement, API, DB, UI, AI worker hoặc infrastructure phải cập nhật change log.
    - Mỗi lần AI/agent chỉnh sửa tài liệu hoặc code liên quan module Workspace phải thêm dòng AI usage log nếu có số liệu usage.
    - Nếu thay đổi DB, phải cập nhật ERD, bảng entity, relationship, index/delete behavior và rollback/cleanup note.
    - Nếu thay đổi API, phải cập nhật route table, DTO/contract notes, happy/unhappy case và web adapter impact.

    ## 2. Review tổng quát

    Module Workspace là lớp tenant boundary Enterprise của WarpTalk. Workspace phân tách dữ liệu, thành viên, meeting, transcript, document library, billing context và policy AI theo từng tổ chức. Code hiện tại chỉ có mô hình Enterprise Workspace; toàn bộ hành vi workspace được điều khiển bằng Owner/Admin/Member, `MembershipType` Internal/External, verified domains và workspace settings. Module hiện được thiết kế theo hướng microservice/Clean Architecture, sử dụng schema `workspace` trong PostgreSQL, Redis cho active context/cache/event stream, gRPC để lấy identity từ Auth Service và để phối hợp policy với TranslationRoom Service.

    Workspace có hai nhóm hành vi chính:

    - **Collaboration boundary:** tạo/chọn Enterprise Workspace, quản lý member, invitation, verified domain và external collaborator.
    - **Knowledge & governance boundary:** quản lý document library, ACL, audit, AI guardrails, encryption/local storage, meeting governance và artifact retention.

    ## 3. Scope và out of scope

    ### In scope

    - Workspace creation, listing, detail, settings và active workspace selection.
    - Enterprise Workspace creation, listing, detail, settings và active workspace selection.
    - Member management, invitation lifecycle, role/ownership rules.
    - Enterprise verified domains và external collaboration.
    - Workspace document library, access policy, audit, AI ingestion/guardrails, local encryption.
    - Workspace governance cho room creation, language policy, artifact access/retention.

    ### Out of scope

    - Non-enterprise workspace type hoặc tự động chuyển đổi giữa các loại workspace.
    - Tự động migrate transcript/artifact giữa các workspace.
    - Custom role nâng cao ngoài Owner/Admin/Member/External Member.
    - Implement code mới; tài liệu này chỉ đặc tả yêu cầu và tổng hợp hiện trạng/spec liên quan.

    ## 4. Kiến trúc và công nghệ sử dụng

    ```mermaid
    {MERMAID["system"].strip()}
    ```

    ### Công nghệ chính

    - **.NET 10 / ASP.NET Core Web API:** REST API và gRPC endpoint của Workspace Service.
    - **Clean Architecture:** API, Application, Domain, Infrastructure tách lớp.
    - **PostgreSQL + EF Core/Npgsql:** lưu schema `workspace`, UUID v7, JSONB settings/policies.
    - **Redis + RabbitMQ:** Redis dùng cho active workspace cache/local stream bridge; RabbitMQ dùng cho durable document/artifact event delivery, publisher confirms, consumer acknowledgements, retry và dead-letter.
    - **gRPC:** lookup identity từ Auth Service, validate workspace/member/policy với TranslationRoom Service.
    - **JWT + signed internal context:** xác thực user và chống spoof workspace context qua downstream headers.
    - **S3/MinIO/Local Storage:** lưu binary document; local provider cần encryption AES-256-CBC + HMAC-SHA512.
    - **AI/RAG/Presidio direction:** AI Service/vector schema xử lý chunk/vector/PII; Workspace chỉ giữ source metadata và guardrail flags.

    ### 4.1 Technology matrix theo subsystem

    | Subsystem | Topic | Technology | Workspace usage |
    |---|---|---|---|
    {tech_table}

    ### 4.2 Test toolchain đã dùng/áp dụng cho backend

    | Tool | Location | Workspace usage |
    |---|---|---|
    {test_toolchain_table}

    ## 5. Database

    ### 5.1 ERD modeling rules applied

    Tài liệu này dùng ERD ở mức **physical data model** vì mục tiêu là phản ánh schema PostgreSQL có thể triển khai từ `warptalk-infrastructure/scripts/init-db.sql`. Theo nguyên tắc ERD phổ biến, entity là bảng, attribute là cột, primary key định danh entity, foreign key thể hiện quan hệ, và cardinality thể hiện một-một/một-nhiều/nhiều-nhiều. Với crow's-foot/Mermaid, `||--o{{` được đọc là một bản ghi phía trái có thể liên kết không hoặc nhiều bản ghi phía phải. Các quan hệ được vẽ theo FK vật lý; với cột workspace_id ở schema khác nhưng không có FK vật lý, tài liệu ghi chú ở boundary/interface thay vì vẽ như FK cứng.

    ```mermaid
    {MERMAID["erd"].strip()}
    ```

    | Entity | Mục đích | Trường chính |
    |---|---|---|
    {db_table}

    ### 5.2 Physical relationship table

    | Parent | Child | Parent cardinality | Child cardinality | FK / behavior |
    |---|---|---:|---:|---|
    {rel_table}

    ### Nguyên tắc database

    - Mọi dữ liệu nghiệp vụ phải scope bởi `workspace_id` khi thuộc workspace.
    - Workspace không hard-delete; record dùng `is_active`, `deleted_at`, `deleted_by`.
    - Document soft-delete/archive không tự động xóa audit trail.
    - Workspace schema không lưu AI chunks; vector/chunk thuộc AI domain.
    - Không cross-join DB giữa Workspace và service khác; dùng gRPC/client boundary.

    ### 5.3 ADR: Quyền tạo meeting theo từng member

    | Decision | Option | Rationale |
    |---|---|---|
    {meeting_creator_decision_table}

    ## 6. API và interface công khai

    | Method | Route | Mục đích |
    |---|---|---|
    {api_table}

    ### Interface nội bộ

    - **Auth gRPC:** resolve user/role snapshot, kiểm tra identity metadata khi list member/invitation.
    - **TranslationRoom gRPC/client:** validate member/policy khi tạo room, join room và xử lý artifact retention.
    - **Redis Stream + RabbitMQ:** Workspace publish document/artifact upload/delete/archive event; Redis giữ local stream/cache bridge, RabbitMQ đảm nhiệm durable async delivery qua exchange/queue/binding, retry và dead-letter cho worker.

    ### 6.3 RabbitMQ messaging workflow

    | Step | Activity | Workspace usage | RabbitMQ rule |
    |---:|---|---|---|
    {rabbitmq_table}
    - **Signed internal context:** Gateway/Auth truyền UserId, ActiveWorkspaceId, Role đã ký cho downstream services.

    ## 7. Web route intent

    | Route | Audience | Intent |
    |---|---|---|
    {web_route_table}

    UI implementation rule: UI phải phân biệt TranslationRoom, MeetingRoom và Workspace Resource; workspace routes cần route guard theo Owner/Admin/Member/External thay vì chỉ kiểm tra token tồn tại. Chi tiết screen behavior, layout, button, loading/empty/error/success state và UI acceptance checklist được tách riêng tại [`workspace-ui-specification.md`](workspace-ui-specification.md), lấy Google Doc UI Mainflow làm source of truth và không dựa vào implementation hiện tại của `warptalk-web`.

    ## 8. Main flow

    ```mermaid
    {MERMAID["main"].strip()}
    ```

    ## 9. Screen flow

    ```mermaid
    {MERMAID["screen"].strip()}
    ```

    ## 10. State diagrams

    ### Invitation lifecycle

    ```mermaid
    {MERMAID["invitation_state"].strip()}
    ```

    ### Document ingestion/access lifecycle

    ```mermaid
    {MERMAID["document_state"].strip()}
    ```

    ### Workspace and Auth eligibility lifecycle

    ```mermaid
    {MERMAID["workspace_account_state"].strip()}
    ```

    ## 11. Use case diagram

    ```mermaid
    {MERMAID["usecase"].strip()}
    ```

    ## 12. User requirements

    | User group | User requirement |
    |---|---|
    | Account user | Có thể tạo, tham gia và chọn active Enterprise Workspace; hệ thống không tự tạo workspace cá nhân mặc định. |
    | Owner | Quản trị settings, domains, invitations, members, roles, ownership, documents, ACL và billing-related policy. |
    | Admin | Quản trị vận hành members/invitations/settings cơ bản và documents theo policy, nhưng không quản lý Owner/billing/delete workspace. |
    | Member | Sử dụng room, transcript, artifact, document theo workspace policy; có thể xem directory khi được phép và rời workspace hợp lệ. |
    | External Member | Chỉ truy cập tài nguyên được mời trực tiếp, không thấy dữ liệu nội bộ hoặc quản trị workspace. |
    | System worker | Ingest/classify documents, publish audit/AI state và tuân thủ policy workspace. |

    ## 13. Functional requirements

    {FUNCTIONAL_REQUIREMENTS_SCOPE}

    | ID | Area | Requirement | Source from specs/code |
    |---|---|---|---|
    {req_table}

    ### 13.1 Functional implementation plan

    | ID | Function | Implementation plan |
    |---|---|---|
    {functional_implementation_plan_table}

    ## 14. Functional test matrix

    | ID | Happy case | Edge case | Unhappy case |
    |---|---|---|---|
    {functional_test_table}

    ### 14.1 Validation and constraint traceability

    | Area | Source | Constraint / validation rule |
    |---|---|---|
    {validation_constraints_table}

    ### 14.2 Existing implemented test coverage

    | Test suite | Scope | Covered cases |
    |---|---|---|
    {existing_test_coverage_table}

    ### 14.3 Predicted additional test cases before/while implementing next scope

    | ID | Case | Condition | Expected result |
    |---|---|---|---|
    {predicted_test_cases_table}

    ## 15. Layer implementation matrix

    | ID | API layer | Application layer | Domain layer | Infrastructure layer |
    |---|---|---|---|---|
    {layer_mapping_table}

    ## 16. Business rules

    {BUSINESS_RULES_SCOPE}

    | ID | Rule | Source from specs/code |
    |---|---|---|
    {br_table}

    ### 16.1 Business rule implementation plan

    | ID | Rule area | Implementation plan |
    |---|---|---|
    {business_rule_implementation_plan_table}

    ## 17. Non-functional requirements

    {NON_FUNCTIONAL_SCOPE}

    | ID | Area | Requirement |
    |---|---|---|
    {nfr_table}

    ## 18. Artifact post-meeting flow

    | Step | Activity | Data affected | Edge/unhappy handling |
    |---|---|---|---|
    {artifact_flow_table}

    ## 19. Happy case / unhappy case

    | ID | Use case | Actor | Happy case | Unhappy case |
    |---|---|---|---|---|
    {uc_table}

    ## 20. Current limitations and cleanup notes

    | Priority | Limitation / cleanup note |
    |---|---|
    {limitation_table}

    ## 21. Future / proposed Workspace scope

    Các mục dưới đây được đưa vào đặc tả trước khi implement để tạo business rule, user story, UI behavior và acceptance criteria rõ ràng. Trạng thái future/proposed không có nghĩa là code hiện tại đã hoàn tất; nó là baseline thiết kế cho implementation tiếp theo.

    | ID | Capability | Description | Status | Source |
    |---|---|---|---|---|
    {future_scope_table}

    ### 21.1 Business rule to user story trace

    | Ticket | Business rules | User story summary | Acceptance source |
    |---|---|---|---|
    {br_user_story_trace_table}

    ## 22. Quality control checklist

    | Area | Checklist item |
    |---|---|
    {qc_table}

    ## 23. Definition of done

    | Area | Done criterion |
    |---|---|
    {dod_table}

    ## 24. Acceptance criteria tổng hợp

    - Tạo workspace thành công luôn tạo Owner membership trong cùng transaction.
    - Hệ thống không expose non-enterprise workspace flows hoặc workspace-type branching.
    - Enterprise Workspace luôn còn ít nhất một Owner active.
    - Invite chỉ accept được khi token hợp lệ và email đăng nhập khớp email được mời.
    - External Member không truy cập workspace settings, directory nội bộ hoặc tài nguyên ngoài meeting scope.
    - Document ACL deny-overrides hoạt động đúng cho explicit deny, sensitive, external và pending ingestion.
    - Document upload sai extension/size hoặc bị policy chặn phải trả lỗi rõ ràng.
    - Redis/RabbitMQ ingestion failure không làm mất document/artifact metadata và phải audit/retry/dead-letter được.
    - TranslationRoom không truy vấn trực tiếp workspace DB; validate qua gRPC/client boundary.
    - Quyền tạo meeting per member phải đọc từ `workspace_members.can_create_meetings`; không dùng JSONB `settings` để lưu danh sách userId allow/deny.
    - Artifact retention được tính từ workspace settings và cleanup không xóa metadata audit bắt buộc.

    ## 25. Traceability nguồn spec

    | Source | Path | Nội dung dùng để tổng hợp |
    |---|---|---|
    {sources_table}
    """)


def font(size: int = 18, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_box(draw, xy, text, fill="#F2F4F7", outline="#A7B3C2", title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    fnt = font(18 if title else 16, bold=title)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 24)
    total_h = len(lines) * 21
    y = y1 + ((y2 - y1 - total_h) // 2)
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill="#0B2545", font=fnt)
        y += 21


def arrow(draw, start, end, color="#41546A"):
    draw.line([start, end], fill=color, width=3)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 12 * direction, ey - 7), (ex - 12 * direction, ey + 7)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 7, ey - 12 * direction), (ex + 7, ey - 12 * direction)]
    draw.polygon(pts, fill=color)


def save_diagram(name: str, title: str, boxes: list[tuple[int, int, int, int, str, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]], size=(1400, 760)):
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 28), title, fill="#1F4D78", font=font(28, bold=True))
    for x1, y1, x2, y2, text, fill in boxes:
        draw_box(draw, (x1, y1, x2, y2), text, fill=fill, title=False)
    for s, e in arrows:
        arrow(draw, s, e)
    path = DIAGRAM_DIR / f"{name}.png"
    img.save(path)
    return path


def build_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {}
    diagrams["system"] = save_diagram(
        "system_context",
        "System / Context Architecture",
        [
            (40, 110, 230, 190, "Client Web/Desktop", "#E8EEF5"),
            (310, 110, 510, 190, "Gateway / Auth", "#E8EEF5"),
            (590, 90, 830, 210, "Workspace Service\nREST + gRPC", "#DDEBF7"),
            (920, 80, 1160, 170, "PostgreSQL\nschema workspace", "#F2F4F7"),
            (920, 200, 1160, 290, "Redis cache\n+ local streams", "#F2F4F7"),
            (1185, 200, 1370, 290, "RabbitMQ\nexchanges + queues", "#F2F4F7"),
            (590, 300, 830, 390, "Auth gRPC\nidentity snapshots", "#F4F6F9"),
            (590, 440, 830, 530, "TranslationRoom\npolicy/artifacts", "#F4F6F9"),
            (920, 410, 1160, 500, "Storage\nS3/MinIO/Local", "#F2F4F7"),
            (920, 560, 1160, 650, "AI/Presidio/RAG\nworker", "#F2F4F7"),
        ],
        [
            ((230, 150), (310, 150)),
            ((510, 150), (590, 150)),
            ((830, 130), (920, 125)),
            ((830, 180), (920, 245)),
            ((705, 210), (705, 300)),
            ((705, 210), (705, 440)),
            ((830, 150), (920, 455)),
            ((1160, 245), (1185, 245)),
            ((1275, 290), (1040, 560)),
        ],
    )
    diagrams["erd"] = save_diagram(
        "erd",
        "Workspace ERD (physical from init-db.sql)",
        [
            (40, 95, 300, 195, "auth.users / auth.roles\nExternal FK reference\nusers: owner/actor\nroles: member/invite role", "#FFF7E6"),
            (470, 120, 760, 230, "workspace.workspaces\nPK id | UK slug\nowner_id -> auth.users\nsettings JSONB, soft delete", "#DDEBF7"),
            (930, 40, 1310, 130, "workspace.workspace_members\nPK id | UK(workspace_id,user_id)\nFK user_id -> auth.users\nFK role_id -> auth.roles\ncan_create_meetings boolean", "#F2F4F7"),
            (930, 160, 1310, 250, "workspace.workspace_invitations\nPK id | UK token_hash\nFK role_id -> auth.roles\nFK invited_by -> auth.users", "#F2F4F7"),
            (930, 280, 1310, 370, "workspace.workspace_verified_domains\nPK id | partial UK(domain) where verified\nFK verified/created/updated_by -> auth.users", "#F2F4F7"),
            (470, 430, 760, 535, "workspace.workspace_documents\nPK id\nFK workspace_id RESTRICT\nAI, sensitivity, retention fields", "#DDEBF7"),
            (930, 420, 1350, 510, "workspace.workspace_document_access_policies\nPK id\nFK document_id CASCADE\nFK workspace_id RESTRICT", "#F4F6F9"),
            (930, 540, 1350, 630, "workspace.workspace_document_audits\nPK id\nFK document_id CASCADE\nFK workspace_id RESTRICT", "#F4F6F9"),
            (930, 660, 1350, 740, "workspace.workspace_knowledge_glossaries\nPK id | UK(workspace,domain,lang,term)\nFK workspace_id RESTRICT", "#F4F6F9"),
            (40, 620, 300, 710, "workspace.schema_migrations\nPK id | UK migration_key\nOperational metadata only", "#F2F4F7"),
        ],
        [
            ((300, 145), (470, 175)),
            ((760, 155), (930, 85)),
            ((760, 175), (930, 205)),
            ((760, 195), (930, 325)),
            ((615, 230), (615, 430)),
            ((760, 482), (930, 465)),
            ((760, 500), (930, 585)),
            ((760, 518), (930, 700)),
        ],
        size=(1420, 790),
    )
    diagrams["main"] = save_diagram(
        "main_flow",
        "Main Flow",
        [
            (60, 140, 250, 220, "Create / Select Workspace", "#DDEBF7"),
            (330, 140, 520, 220, "Invite / Manage Members", "#E8EEF5"),
            (600, 140, 790, 220, "Enterprise Domain Policy", "#FFF7E6"),
            (870, 80, 1080, 160, "Internal Role", "#F2F4F7"),
            (870, 220, 1080, 300, "External Member", "#F2F4F7"),
            (1160, 140, 1360, 220, "Room / Document Action", "#DDEBF7"),
            (600, 400, 790, 480, "Policy Check", "#FFF7E6"),
            (870, 400, 1080, 480, "Persist + Publish", "#F2F4F7"),
            (1160, 400, 1360, 480, "Audit / Retention / AI", "#F2F4F7"),
        ],
        [
            ((250, 180), (330, 180)),
            ((520, 180), (600, 180)),
            ((790, 160), (870, 120)),
            ((790, 200), (870, 260)),
            ((1080, 120), (1160, 170)),
            ((1080, 260), (1160, 190)),
            ((1260, 220), (695, 400)),
            ((790, 440), (870, 440)),
            ((1080, 440), (1160, 440)),
        ],
    )
    diagrams["screen"] = save_diagram(
        "screen_flow",
        "Screen Flow",
        [
            (60, 120, 230, 190, "Login", "#E8EEF5"),
            (310, 120, 520, 190, "Workspace Switcher", "#DDEBF7"),
            (600, 120, 820, 190, "Dashboard", "#DDEBF7"),
            (900, 40, 1110, 110, "Members", "#F2F4F7"),
            (900, 135, 1110, 205, "Invitations", "#F2F4F7"),
            (900, 230, 1110, 300, "Documents", "#F2F4F7"),
            (900, 325, 1110, 395, "Settings", "#F2F4F7"),
            (1170, 230, 1350, 300, "Document Detail / ACL", "#F4F6F9"),
            (1170, 325, 1350, 395, "Verified Domains", "#F4F6F9"),
        ],
        [
            ((230, 155), (310, 155)),
            ((520, 155), (600, 155)),
            ((820, 155), (900, 75)),
            ((820, 155), (900, 170)),
            ((820, 155), (900, 265)),
            ((820, 155), (900, 360)),
            ((1110, 265), (1170, 265)),
            ((1110, 360), (1170, 360)),
        ],
    )
    diagrams["invitation_state"] = save_diagram(
        "invitation_state",
        "Invitation State Diagram",
        [
            (90, 250, 250, 320, "Pending", "#DDEBF7"),
            (430, 80, 600, 150, "Accepted", "#E7F4E4"),
            (430, 210, 600, 280, "Revoked", "#F9E3E3"),
            (430, 340, 600, 410, "Expired", "#F9E3E3"),
            (430, 470, 600, 540, "Replaced", "#FFF7E6"),
        ],
        [
            ((250, 285), (430, 115)),
            ((250, 285), (430, 245)),
            ((250, 285), (430, 375)),
            ((250, 285), (430, 505)),
        ],
        size=(760, 620),
    )
    diagrams["document_state"] = save_diagram(
        "document_state",
        "Document Ingestion / Access State",
        [
            (60, 110, 230, 180, "Uploaded", "#E8EEF5"),
            (310, 110, 520, 180, "Pending Ingestion", "#FFF7E6"),
            (600, 60, 820, 130, "Awaiting Approval", "#FFF7E6"),
            (600, 185, 820, 255, "Completed", "#E7F4E4"),
            (900, 60, 1070, 130, "Rejected", "#F9E3E3"),
            (900, 185, 1070, 255, "Archived", "#F2F4F7"),
            (1150, 125, 1320, 195, "Deleted", "#F9E3E3"),
        ],
        [
            ((230, 145), (310, 145)),
            ((520, 145), (600, 95)),
            ((520, 145), (600, 220)),
            ((820, 95), (900, 95)),
            ((820, 220), (900, 220)),
            ((1070, 95), (1150, 145)),
            ((1070, 220), (1150, 160)),
        ],
        size=(1400, 360),
    )
    diagrams["workspace_account_state"] = save_diagram(
        "workspace_account_state",
        "Workspace / Auth Eligibility State",
        [
            (60, 120, 250, 190, "Active Workspace", "#E7F4E4"),
            (330, 55, 540, 125, "Inactive Workspace", "#FFF7E6"),
            (330, 190, 540, 260, "Soft Deleted Workspace", "#F9E3E3"),
            (620, 120, 810, 190, "Member Removed", "#FFF7E6"),
            (890, 55, 1120, 125, "User Still ACTIVE", "#E7F4E4"),
            (890, 190, 1190, 260, "Auth SUSPENDED_NO_ACTIVE_WORKSPACE", "#F9E3E3"),
            (1230, 190, 1390, 260, "Accept Active Invite", "#DDEBF7"),
        ],
        [
            ((250, 155), (330, 90)),
            ((250, 155), (330, 225)),
            ((540, 90), (60, 155)),
            ((250, 155), (620, 155)),
            ((810, 155), (890, 90)),
            ((810, 155), (890, 225)),
            ((540, 90), (890, 225)),
            ((540, 225), (890, 225)),
            ((1190, 225), (1230, 225)),
            ((1390, 225), (60, 155)),
        ],
        size=(1450, 340),
    )
    diagrams["usecase"] = save_diagram(
        "usecase",
        "Use Case Diagram",
        [
            (60, 95, 230, 165, "Owner", "#DDEBF7"),
            (60, 205, 230, 275, "Admin", "#DDEBF7"),
            (60, 315, 230, 385, "Member", "#DDEBF7"),
            (60, 425, 230, 495, "External Member", "#DDEBF7"),
            (60, 535, 230, 605, "System Worker", "#DDEBF7"),
            (420, 60, 700, 130, "Create/select workspace", "#F2F4F7"),
            (420, 155, 700, 225, "Manage settings/domains", "#F2F4F7"),
            (420, 250, 700, 320, "Invite/manage members", "#F2F4F7"),
            (420, 345, 700, 415, "Manage documents/ACL", "#F2F4F7"),
            (420, 440, 700, 510, "Use rooms/documents by policy", "#F2F4F7"),
            (420, 535, 700, 605, "Ingest/classify/audit", "#F2F4F7"),
        ],
        [
            ((230, 130), (420, 95)),
            ((230, 130), (420, 190)),
            ((230, 130), (420, 285)),
            ((230, 130), (420, 380)),
            ((230, 240), (420, 285)),
            ((230, 240), (420, 380)),
            ((230, 350), (420, 475)),
            ((230, 460), (420, 475)),
            ((230, 570), (420, 570)),
        ],
        size=(820, 680),
    )
    return diagrams


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    set_table_borders(table)


def add_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], LIGHT)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_diagram(doc: Document, title: str, image_path: Path, caption: str):
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(95, 95, 95)


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "WarpTalk - Workspace Module SRS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(95, 95, 95)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - Capstone project documentation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(95, 95, 95)


def add_cover(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Software Requirement Specification")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Module Workspace - WarpTalk Backend")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    meta_rows = [
        ("Ngôn ngữ", "Tiếng Việt"),
        ("Phiên bản", "1.6"),
        ("Người tạo file", "Ngô Xuân Hạnh Nhi"),
        ("Ngày", "2026-06-11"),
        ("Phạm vi", "Workspace Service, Auth/TranslationRoom integration, Document governance"),
        ("Nguồn", "Tổng hợp từ specs workspace và kiểm chứng code hiện tại"),
    ]
    doc.add_paragraph()
    add_table(doc, ["Thuộc tính", "Giá trị"], meta_rows, [2600, 6400])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tài liệu này không thay đổi code backend; mục tiêu là chuẩn hóa requirement và kiến trúc module Workspace.").italic = True
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_docx(diagrams: dict[str, Path]):
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Tài liệu SRS này mô tả yêu cầu phần mềm cho module Workspace của WarpTalk Backend. "
        "Module Workspace đóng vai trò tenant boundary, quản lý không gian làm việc, thành viên, lời mời, verified domain, document library, AI guardrails và chính sách governance cho meeting/artifact."
    )
    add_table(
        doc,
        ["Mục", "Mô tả"],
        [
            ("Purpose", "Chuẩn hóa requirement, rule, flow, database và interface cho module Workspace."),
            ("Audience", "Backend engineers, QA, reviewer kiến trúc, capstone evaluator, product owner."),
            ("Scope", "Workspace Service và các tích hợp Auth, TranslationRoom, Transcript, AI/Document ingestion."),
            ("Non-goal", "Không triển khai code mới, không thay đổi API, không tạo migration."),
        ],
        [2100, 7200],
    )

    doc.add_heading("1.1 Document Control", level=2)
    add_table(doc, ["Field", "Value"], DOC_CONTROL_ROWS, [2400, 6900])
    doc.add_heading("1.2 Change Log", level=2)
    add_table(doc, ["Version", "Date", "Author/AI", "Change", "Reason"], CHANGE_LOG, [900, 1400, 1400, 2700, 2900])
    doc.add_heading("1.3 AI Usage Log", level=2)
    add_table(doc, ["Date", "AI/Actor", "Scope", "Work performed", "Usage"], AI_USAGE_LOG, [1300, 1400, 1700, 3300, 1600])
    doc.add_heading("1.4 Rules For Updating This File", level=2)
    add_bullets(
        doc,
        [
            "Every material edit must update changelog, source references and QA checklist.",
            "Any AI/agent change to workspace docs or code must add an AI usage log row when telemetry is available.",
            "DB changes must update ERD, entity table, relationship table, index/delete behavior and rollback notes.",
            "API changes must update route table, DTO/contract notes, happy/unhappy cases and web adapter impact.",
        ],
    )

    doc.add_heading("2. Overall Description", level=1)
    doc.add_paragraph(
        "Workspace là đơn vị cô lập dữ liệu và chính sách Enterprise trong WarpTalk. Code hiện tại chỉ có mô hình Enterprise Workspace. "
        "Người dùng có thể tạo, tham gia và chọn active Enterprise Workspace; "
        "hành vi internal/external được kiểm soát bằng verified domain, membership type và external collaboration policy."
    )
    add_bullets(
        doc,
        [
            "Enterprise Workspace: tenant boundary duy nhất, nhiều thành viên, role Owner/Admin/Member.",
            "MembershipType: Internal hoặc External, dùng để kiểm soát directory, invitation, domain và resource visibility.",
            "Verified domains: dùng cho internal membership enforcement và chống nhầm lẫn tenant.",
            "Document governance: document metadata, ACL, audit, AI ingestion, local encryption, retention.",
            "Meeting governance: policy kiểm soát room creation, language, external access và artifact retention.",
        ],
    )

    doc.add_heading("3. System Architecture", level=1)
    add_diagram(doc, "3.1 System Context", diagrams["system"], "Workspace Service phối hợp với Auth, TranslationRoom, PostgreSQL, Redis, Storage và AI workers.")
    doc.add_paragraph(
        "Service boundary chính: Workspace không đọc database của Auth hoặc TranslationRoom. Identity enrichment dùng Auth gRPC client; meeting validation và artifact policy dùng gRPC/client boundary với TranslationRoom. "
        "Active workspace context được lưu trong Redis/session và được Gateway/Auth ký khi truyền cho downstream services."
    )
    add_table(
        doc,
        ["Layer / Component", "Responsibility"],
        [
            ("API", "REST controllers cho workspace, members, invitations, documents; gRPC invitation service."),
            ("Application", "Business rules, validation, DTO mapping, document access evaluator, service orchestration."),
            ("Domain", "Entities, enums, constants, settings, role/membership vocabulary."),
            ("Infrastructure", "EF Core repositories, PostgreSQL DbContext, Redis cache/stream, gRPC clients, background ingestion worker."),
            ("Storage/AI", "Binary file storage và downstream AI/RAG/PII processing ngoài workspace domain."),
        ],
        [2600, 6600],
    )

    doc.add_heading("4. Technology Stack", level=1)
    add_table(
        doc,
        ["Technology", "Usage"],
        [
            (".NET 10 / ASP.NET Core", "REST API, DI, authentication, gRPC hosting/client."),
            ("EF Core + Npgsql", "Mapping schema workspace, UUID v7, PostgreSQL indexes, JSONB settings."),
            ("PostgreSQL", "Persistent workspace/member/invitation/document metadata and audit trail."),
            ("Redis + RabbitMQ", "Redis for active workspace cache/local stream bridge; RabbitMQ for durable document/artifact ingestion events, publisher confirms, consumer acknowledgements, retry and dead-letter."),
            ("gRPC", "Typed inter-service communication with Auth and TranslationRoom services."),
            ("JWT + signed internal context", "Authentication and tamper-resistant downstream workspace context."),
            ("S3/MinIO/Local Storage", "Document binary storage; local provider requires application-level encryption."),
            ("AES-256-CBC + HMAC-SHA512", "Encrypt-then-MAC for local document storage per workspace-derived key."),
        ],
        [2700, 6500],
    )
    doc.add_heading("4.1 Workspace Technology Matrix", level=2)
    add_table(doc, ["Subsystem", "Topic", "Technology", "Workspace usage"], TECH_MATRIX, [1500, 1700, 2300, 3800])
    doc.add_heading("4.2 Backend Test Toolchain", level=2)
    add_table(doc, ["Tool", "Location", "Workspace usage"], TEST_TOOLCHAIN, [2200, 3200, 3800])

    doc.add_heading("5. Data Requirements", level=1)
    doc.add_heading("5.1 ERD Modeling Rules Applied", level=2)
    doc.add_paragraph(
        "This SRS uses a physical ERD for the Workspace schema because the goal is to reflect deployable PostgreSQL tables from warptalk-infrastructure/scripts/init-db.sql. "
        "Entities map to tables, attributes map to columns, PK/UK/FK constraints are shown, and cardinality follows crow's-foot semantics. "
        "Relationship lines are drawn only for physical foreign keys; cross-service workspace_id references without FK are documented as integration boundaries."
    )
    add_diagram(doc, "5.1 Physical ERD", diagrams["erd"], "Physical ERD của schema workspace sinh theo warptalk-infrastructure/scripts/init-db.sql.")
    add_table(doc, ["Entity", "Purpose", "Key fields"], DB_ENTITIES, [2800, 2600, 3900])
    doc.add_heading("5.2 Physical Relationships", level=2)
    add_table(doc, ["Parent", "Child", "Parent", "Child", "FK / behavior"], PHYSICAL_RELATIONSHIPS, [2100, 2900, 900, 900, 2500])
    doc.add_heading("5.3 ADR: Per-member Meeting Creator Permission", level=2)
    doc.add_paragraph(
        "Quyền tạo meeting là per-membership permission. Workspace chọn cột `workspace_members.can_create_meetings` thay vì lưu danh sách userId allow/deny trong `workspace.settings` JSONB."
    )
    add_table(doc, ["Decision", "Option", "Rationale"], MEETING_CREATOR_PERMISSION_DECISION, [1200, 2800, 5300])

    doc.add_heading("6. External Interface Requirements", level=1)
    doc.add_heading("6.1 REST APIs", level=2)
    add_table(doc, ["Method", "Route", "Purpose"], API_ROWS, [1300, 4200, 3800])
    doc.add_heading("6.2 Internal Interfaces", level=2)
    add_bullets(
        doc,
        [
            "Auth gRPC resolves user/role snapshots; Workspace keeps membership role assignment, not Auth navigation entities.",
            "TranslationRoom integration validates workspace member, CanCreateMeetings, allowed target languages and artifact retention policy.",
            "Redis Stream plus RabbitMQ publishes document/artifact upload/delete/archive events for AI ingestion, embedding invalidation, retry and dead-letter.",
            "Signed internal context carries UserId, ActiveWorkspaceId and role across service hops.",
        ],
    )

    doc.add_heading("6.3 RabbitMQ Messaging Workflow", level=2)
    add_table(doc, ["Step", "Activity", "Workspace usage", "RabbitMQ rule"], RABBITMQ_WORKFLOW, [700, 1700, 3500, 3400])

    doc.add_heading("6.4 Web Route Intent", level=2)
    add_table(doc, ["Route", "Audience", "Intent"], WEB_ROUTES, [2700, 2400, 4100])
    doc.add_paragraph(
        "UI must distinguish TranslationRoom, MeetingRoom and Workspace Resource. Workspace routes require role-aware guards for Owner/Admin/Member/External Member, not only access-token presence. Detailed screen behavior is separated into workspace-ui-specification.md and uses the UI Mainflow Google Doc as source of truth, not the current warptalk-web implementation."
    )

    doc.add_heading("7. Functional Requirements", level=1)
    doc.add_paragraph(FUNCTIONAL_REQUIREMENTS_SCOPE)
    add_table(doc, ["ID", "Area", "Requirement", "Source from specs/code"], functional_requirement_rows(), [1000, 1300, 3900, 3100])
    doc.add_heading("7.1 Functional Implementation Plan", level=2)
    add_table(doc, ["ID", "Function", "Implementation plan"], FUNCTIONAL_IMPLEMENTATION_PLAN, [1000, 1900, 6400])

    doc.add_heading("8. Functional Test Matrix", level=1)
    add_table(doc, ["ID", "Happy case", "Edge case", "Unhappy case"], FUNCTIONAL_TEST_CASES, [1000, 2800, 2700, 2800])
    doc.add_heading("8.1 Validation And Constraint Traceability", level=2)
    add_table(doc, ["Area", "Source", "Constraint / validation rule"], VALIDATION_CONSTRAINTS, [1900, 3000, 4300])
    doc.add_heading("8.2 Existing Implemented Test Coverage", level=2)
    add_table(doc, ["Test suite", "Scope", "Covered cases"], EXISTING_TEST_COVERAGE, [2400, 2300, 4600])
    doc.add_heading("8.3 Predicted Additional Test Cases", level=2)
    add_table(doc, ["ID", "Case", "Condition", "Expected result"], PREDICTED_TEST_CASES, [1200, 2300, 3000, 2800])

    doc.add_heading("9. Layer Implementation Matrix", level=1)
    add_table(doc, ["ID", "API", "Application", "Domain", "Infrastructure"], FUNCTIONAL_LAYER_MAPPING, [900, 2100, 2300, 1900, 2100])

    doc.add_heading("10. Business Rules", level=1)
    doc.add_paragraph(BUSINESS_RULES_SCOPE)
    add_table(doc, ["ID", "Rule", "Source from specs/code"], business_rule_rows(), [1100, 5200, 3000])
    doc.add_heading("10.1 Business Rule Implementation Plan", level=2)
    add_table(doc, ["ID", "Rule area", "Implementation plan"], BUSINESS_RULE_IMPLEMENTATION_PLAN, [1100, 2100, 6100])

    doc.add_heading("11. Non-functional Requirements", level=1)
    doc.add_paragraph(NON_FUNCTIONAL_SCOPE)
    add_table(doc, ["ID", "Area", "Requirement"], NFRS, [1400, 1700, 6100])

    doc.add_heading("12. User Requirements and Use Cases", level=1)
    add_diagram(doc, "12.1 Use Case Diagram", diagrams["usecase"], "Actor và capability chính của module Workspace.")
    add_table(
        doc,
        ["User group", "Requirement"],
        [
            ("Account user", "Có thể tạo, tham gia và chọn active Enterprise Workspace; hệ thống không tự tạo workspace cá nhân mặc định."),
            ("Owner", "Quản trị settings, domains, invitations, members, roles, ownership, documents, ACL và policy."),
            ("Admin", "Quản trị vận hành members/invitations/settings cơ bản và documents theo policy."),
            ("Member", "Sử dụng room, transcript, artifact, document theo workspace policy và có thể rời workspace hợp lệ."),
            ("External Member", "Chỉ truy cập tài nguyên được mời trực tiếp; không thấy dữ liệu nội bộ."),
            ("System worker", "Ingest/classify documents, cập nhật trạng thái, audit và AI guardrails."),
        ],
        [2300, 7000],
    )
    add_table(
        doc,
        ["ID", "Use case", "Actor", "Precondition", "Postcondition"],
        [(uid, name, actor, pre, post) for uid, name, actor, pre, happy, post, unhappy in USE_CASES],
        [1000, 2100, 1700, 2600, 2100],
    )

    doc.add_heading("13. Main Flow, Screen Flow and State Models", level=1)
    add_diagram(doc, "13.1 Main Flow", diagrams["main"], "Flow tổng quát từ workspace context tới member, policy, room/document và audit/AI.")
    add_diagram(doc, "13.2 Screen Flow", diagrams["screen"], "Screen flow tham chiếu cho workspace switcher, dashboard, members, invitations, documents và settings.")
    add_diagram(doc, "13.3 Invitation State Diagram", diagrams["invitation_state"], "Invitation token lifecycle.")
    add_diagram(doc, "13.4 Document State Diagram", diagrams["document_state"], "Document lifecycle từ upload tới ingestion, approval, archive/delete.")
    add_diagram(doc, "13.5 Workspace/Auth Eligibility State Diagram", diagrams["workspace_account_state"], "Workspace deactivation/member removal invalidates context and can suspend Auth app eligibility when no active workspace remains.")

    doc.add_heading("14. Artifact Post-meeting Flow", level=1)
    add_table(doc, ["Step", "Activity", "Data affected", "Edge/unhappy handling"], ARTIFACT_POST_MEETING_FLOW, [900, 2600, 3000, 2800])

    doc.add_heading("15. Happy Case / Unhappy Case", level=1)
    add_table(
        doc,
        ["ID", "Use case", "Happy case", "Unhappy case"],
        [(uid, name, happy, unhappy) for uid, name, actor, pre, happy, post, unhappy in USE_CASES],
        [900, 2000, 3300, 3100],
    )

    doc.add_heading("16. Acceptance Criteria", level=1)
    add_bullets(
        doc,
        [
            "Workspace creation always creates Owner membership atomically.",
            "System does not expose non-enterprise workspace flows or workspace-type branching.",
            "Enterprise Workspace always retains at least one active Owner.",
            "Invitation accept requires valid token and exact email match.",
            "External Member cannot access workspace settings, internal directory or resources outside direct meeting scope.",
            "Document ACL deny-overrides correctly handles explicit deny, sensitive, external and pending ingestion cases.",
            "Document upload rejects unsupported file type/size and records clear errors.",
            "Redis/RabbitMQ ingestion failure does not lose document/artifact metadata and remains retry/audit/dead-letter friendly.",
            "TranslationRoom validates workspace policy through gRPC/client boundary, never direct workspace DB cross-join.",
            "Artifact retention follows workspace settings while preserving required audit metadata.",
        ],
    )

    doc.add_heading("17. Current Limitations And Cleanup Notes", level=1)
    add_table(doc, ["Priority", "Limitation / cleanup note"], LIMITATIONS, [1300, 8000])

    doc.add_heading("18. Future / Proposed Workspace Scope", level=1)
    doc.add_paragraph(
        "Các mục dưới đây được đưa vào đặc tả trước khi implement để tạo business rule, user story, UI behavior và acceptance criteria rõ ràng. "
        "Trạng thái future/proposed không có nghĩa là code hiện tại đã hoàn tất; nó là baseline thiết kế cho implementation tiếp theo."
    )
    add_table(doc, ["ID", "Capability", "Description", "Status", "Source"], FUTURE_PROPOSED_SCOPE, [900, 1700, 3600, 1700, 1400])

    doc.add_heading("18.1 Business Rule To User Story Trace", level=2)
    add_table(doc, ["Ticket", "Business rules", "User story summary", "Acceptance source"], BUSINESS_RULE_USER_STORY_TRACE, [1000, 1800, 3400, 3100])

    doc.add_heading("19. Quality Control Checklist", level=1)
    add_table(doc, ["Area", "Checklist item"], QC_CHECKLIST, [1800, 7400])

    doc.add_heading("20. Definition Of Done", level=1)
    add_table(doc, ["Area", "Done criterion"], DEFINITION_OF_DONE, [1800, 7400])

    doc.add_heading("21. Source Traceability", level=1)
    add_table(doc, ["Source", "Path", "Used content"], SOURCE_SPECS, [1700, 4200, 3400])

    doc.add_heading("22. Assumptions and Defaults", level=1)
    add_bullets(
        doc,
        [
            "Tài liệu là SRS/overview, không phải migration hoặc implementation patch.",
            "Các đoạn spec nguồn bị lỗi encoding được diễn giải lại bằng tiếng Việt đúng nghĩa kỹ thuật.",
            "Diagram trong DOCX là ảnh tĩnh để không phụ thuộc Mermaid runtime trong Microsoft Word.",
            "Mermaid source vẫn được giữ trong Markdown để dễ review/diff trong repo.",
            "Nếu render DOCX không khả dụng do thiếu LibreOffice, artifact DOCX vẫn được tạo và báo rõ trong kết quả.",
        ],
    )

    doc.save(DOCX_PATH)


def main():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(make_markdown(), encoding="utf-8")
    diagrams = build_diagrams()
    build_docx(diagrams)
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote diagrams to {DIAGRAM_DIR}")


if __name__ == "__main__":
    main()
